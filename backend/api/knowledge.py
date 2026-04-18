"""
知识管理 API — PostgreSQL 持久化
支持文档上传、列表查询、删除、搜索
"""

import time
import uuid
import tiktoken
import re
import json
import tempfile
from datetime import datetime
from collections import Counter
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, HTTPException, Form, Depends, BackgroundTasks
from fastapi.responses import FileResponse
import aiofiles
from pydantic import BaseModel
from docx import Document
from backend.api.auth import require_expert, get_current_user

from backend.core.parser.document import parse_document
from backend.core.rag.pgvector_retriever import add_chunks_to_db, retrieve
from backend.core.llm.manager import get_llm_manager
from backend.core.knowledge.ai_graph_extractor import extract_knowledge_items_with_ai
from backend.models.schemas import UploadResponse
from backend.config import settings
import psycopg2, psycopg2.extras

router = APIRouter(tags=["知识管理"])
DEVICE_TERMS = [
    "手持式无线吸尘器", "无线吸尘器", "吸尘器",
    "电饭煲", "传送带", "输送带",
    "电机", "液压泵", "泵", "阀门", "传感器", "轴承", "PLC", "变频器",
    "液压缸", "输送带", "风机", "压缩机", "电池", "继电器", "接触器",
    "控制器", "电源", "减速机", "编码器", "过滤器"
]
CANONICAL_DEVICE_ALIASES = [
    ("吸尘器", ["手持式无线吸尘器", "无线吸尘器", "吸尘器"]),
    ("电饭煲", ["电饭煲", "电饭锅"]),
    ("传送带", ["传送带", "输送带"]),
]
FAULT_HINTS = [
    "故障", "异常", "报警", "失效", "损坏", "泄漏", "过热", "振动", "异响",
    "堵塞", "磨损", "卡滞", "偏差", "无法启动", "不启动", "无压力", "压力不足",
    "温度过高", "短路", "断路", "跳闸", "停机"
]
SOLUTION_HINTS = [
    "检查", "更换", "清理", "维修", "修复", "调整", "校准", "紧固",
    "润滑", "复位", "重启", "测试", "确认", "处理", "排查",
    "连接", "按压", "激活", "检测", "冲洗", "清空", "取出", "滴加", "清洁", "冷却", "待",
    "拆卸", "测量", "插拔", "更换", "送修", "购买", "装回"
]
COMMON_FAULT_TITLES = [
    "无法开机", "吸力减弱", "异常噪音", "充电故障",
    "无法启动", "不启动", "无法充电", "充不进电",
    "过热报警", "异响", "无压力", "压力不足"
]
NOISE_TERMS = {
    "常见故障排查表", "技术参数", "定期保养计划", "安全须知", "产品结构图示", "分步维修指南",
    "每日", "每周", "每月", "每半年", "每年", "额定电压", "空载转速", "最大真空度",
    "电池容量", "工作噪音", "注意事项", "维修查询热线", "更新日期", "可能原因", "解决方法",
    "故障现象", "本手册", "标准化模板", "官方完整手册", "维修保养手册", "常见故障",
    "故障诊断", "电路图", "流程图", "目录"
}


def _normalize_pipeline(pipeline: str | None) -> str:
    p = (pipeline or "").strip()
    if not p:
        return "流水线1"
    if len(p) > 64:
        raise HTTPException(status_code=400, detail="流水线名称过长（最多64字符）")
    return p


def _ensure_structured_knowledge_tables(conn) -> None:
    with conn.cursor() as cur:
        cur.execute("CREATE EXTENSION IF NOT EXISTS vector")
        cur.execute("CREATE EXTENSION IF NOT EXISTS pgcrypto")
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS knowledge_items (
                item_id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                pipeline VARCHAR(64) NOT NULL DEFAULT '流水线1',
                machine_category VARCHAR(120) NOT NULL DEFAULT '',
                machine VARCHAR(160) NOT NULL DEFAULT '',
                problem_category VARCHAR(120) NOT NULL DEFAULT '',
                problem TEXT NOT NULL,
                root_cause TEXT NOT NULL DEFAULT '',
                solution TEXT NOT NULL DEFAULT '',
                metadata JSONB NOT NULL DEFAULT '{}'::jsonb,
                status VARCHAR(20) NOT NULL DEFAULT 'active',
                created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
            )
            """
        )
        cur.execute(
            """
            CREATE INDEX IF NOT EXISTS idx_knowledge_items_pipeline
            ON knowledge_items(pipeline)
            """
        )
        cur.execute(
            """
            WITH ranked AS (
                SELECT
                    ctid,
                    row_number() OVER (
                        PARTITION BY pipeline, machine, problem, root_cause
                        ORDER BY updated_at DESC, created_at DESC
                    ) AS rn
                FROM knowledge_items
            )
            DELETE FROM knowledge_items k
            USING ranked r
            WHERE k.ctid = r.ctid
              AND r.rn > 1
            """
        )
        cur.execute(
            """
            CREATE UNIQUE INDEX IF NOT EXISTS uq_knowledge_items_unique
            ON knowledge_items (pipeline, machine, problem, root_cause)
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS knowledge_item_weights (
                item_id UUID PRIMARY KEY REFERENCES knowledge_items(item_id) ON DELETE CASCADE,
                helpful_weight DOUBLE PRECISION NOT NULL DEFAULT 0,
                misleading_weight DOUBLE PRECISION NOT NULL DEFAULT 0,
                feedback_count INTEGER NOT NULL DEFAULT 0,
                current_weight DOUBLE PRECISION NOT NULL DEFAULT 0.5,
                expert_weight DOUBLE PRECISION,
                updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
            )
            """
        )
        cur.execute("ALTER TABLE knowledge_item_weights ADD COLUMN IF NOT EXISTS expert_weight DOUBLE PRECISION")
        conn.commit()


def _ensure_pipeline_registry_table(conn) -> None:
    with conn.cursor() as cur:
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS pipeline_registry (
                pipeline VARCHAR(64) PRIMARY KEY,
                created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
            )
            """
        )
        cur.execute(
            """
            INSERT INTO pipeline_registry(pipeline)
            VALUES ('流水线1')
            ON CONFLICT (pipeline) DO NOTHING
            """
        )
        conn.commit()


class KnowledgeItemCreateRequest(BaseModel):
    pipeline: str = "流水线1"
    machine_category: str = ""
    machine: str = ""
    problem_category: str = ""
    problem: str
    root_cause: str = ""
    solution: str = ""
    metadata: dict | None = None


def _clean_single_line_text(raw: str, max_len: int) -> str:
    s = str(raw or "").replace("\r", " ").replace("\n", " ").replace("\t", " ")
    s = re.sub(r"\s+", " ", s).strip()
    if not s:
        return ""
    return s[:max_len].strip()


def _validate_knowledge_item_fields(pipeline: str, machine_category: str, machine: str, problem_category: str, problem: str, root_cause: str, solution: str):
    if not pipeline:
        raise HTTPException(status_code=400, detail="pipeline 不能为空")
    if not problem:
        raise HTTPException(status_code=400, detail="problem 不能为空")
    if not root_cause:
        raise HTTPException(status_code=400, detail="root_cause 不能为空")
    if len(problem) > 240:
        raise HTTPException(status_code=400, detail="problem 太长（最多240字）")
    if len(root_cause) > 320:
        raise HTTPException(status_code=400, detail="root_cause 太长（最多320字）")
    if len(machine) > 160:
        raise HTTPException(status_code=400, detail="machine 太长（最多160字）")
    if len(machine_category) > 120:
        raise HTTPException(status_code=400, detail="machine_category 太长（最多120字）")
    if len(problem_category) > 120:
        raise HTTPException(status_code=400, detail="problem_category 太长（最多120字）")
    if len(solution) > 800:
        raise HTTPException(status_code=400, detail="solution 太长（最多800字）")
    if any(k in problem for k in ["来源：", "来源:", "doc_id", "chunk_id"]):
        raise HTTPException(status_code=400, detail="problem 包含不允许的来源信息")
    if any(k in root_cause for k in ["来源：", "来源:", "doc_id", "chunk_id"]):
        raise HTTPException(status_code=400, detail="root_cause 包含不允许的来源信息")


async def _ai_enrich_categories(pipeline: str, machine: str, problem: str, root_cause: str) -> dict:
    prompt = f"""
你是工业知识库结构化助手。请根据输入内容，为知识库条目补全分类字段。

要求：
1) 严格只输出 JSON，不要任何解释文字
2) machine_category 与 problem_category 必须尽量简短（1~6个字），若无法判断返回空字符串
3) problem_category 必须优先从这些大类中选择：电气/机械/控制/液压/气动/传感器/环境/维护/使用/配件/性能/其他
3) 不要输出来源/文件名/页码/doc_id 等信息

输出 JSON 结构：
{{
  "machine_category": "例如 电机/变频器/PLC/传感器/液压/气动/轴承/传送带/其他 或 空",
  "problem_category": "例如 电气/机械/控制/液压/气动/传感器/环境/维护/使用/配件/性能/其他 或 空"
}}

流水线：{pipeline}
设备：{machine}
现象：{problem}
根因：{root_cause}
"""
    manager = get_llm_manager()
    resp, _provider = await manager.generate_with_fallback(prompt)
    parsed = _extract_json_dict(resp.content if resp else "")
    if not isinstance(parsed, dict):
        return {}
    out = {}
    out["machine_category"] = _clean_single_line_text(parsed.get("machine_category", ""), 120)
    out["problem_category"] = _clean_single_line_text(parsed.get("problem_category", ""), 120)
    return out


async def _ai_fill_item_blanks(pipeline: str, machine_category: str, machine: str, problem_category: str, problem: str, root_cause: str, solution: str, metadata: dict) -> dict:
    meta = metadata if isinstance(metadata, dict) else {}
    meta_text = json.dumps(meta, ensure_ascii=False)
    if len(meta_text) > 1200:
        meta_text = meta_text[:1200]
    prompt = f"""
你是工业知识库条目补全助手。请根据已知信息，补全缺失字段。

要求：
1) 严格只输出 JSON，不要任何解释文字
2) 字段尽量简洁、可落地
3) 不要输出来源/文件名/页码/doc_id/chunk_id 等信息
4) 对于无法判断的字段，返回空字符串

输出 JSON 结构：
{{
  "machine_category": "",
  "machine": "",
  "problem_category": "",
  "problem": "",
  "root_cause": "",
  "solution": ""
}}

流水线：{pipeline}
machine_category：{machine_category}
machine：{machine}
problem_category：{problem_category}
problem：{problem}
root_cause：{root_cause}
solution：{solution}
metadata：{meta_text}
"""
    manager = get_llm_manager()
    resp, _provider = await manager.generate_with_fallback(prompt)
    parsed = _extract_json_dict(resp.content if resp else "")
    if not isinstance(parsed, dict):
        return {}
    out = {}
    out["machine_category"] = _clean_single_line_text(parsed.get("machine_category", ""), 120)
    out["machine"] = _clean_single_line_text(parsed.get("machine", ""), 160)
    out["problem_category"] = _clean_single_line_text(parsed.get("problem_category", ""), 120)
    out["problem"] = _clean_single_line_text(parsed.get("problem", ""), 240)
    out["root_cause"] = _clean_single_line_text(parsed.get("root_cause", ""), 320)
    out["solution"] = _clean_single_line_text(parsed.get("solution", ""), 800)
    return out


class KnowledgeItemUpdateRequest(BaseModel):
    pipeline: str | None = None
    machine_category: str | None = None
    machine: str | None = None
    problem_category: str | None = None
    problem: str | None = None
    root_cause: str | None = None
    solution: str | None = None
    metadata: dict | None = None
    status: str | None = None


class KnowledgeItemSearchRequest(BaseModel):
    query: str = ""
    pipeline: str | None = None
    top_k: int = 10


class KnowledgeItemWeightFeedbackRequest(BaseModel):
    item_id: str
    feedback_type: str
    amount: float = 1.0


class KnowledgeItemExpertWeightRequest(BaseModel):
    item_id: str
    expert_weight: float | None = None


class KnowledgeItemReextractRequest(BaseModel):
    pipeline: str = "流水线1"
    doc_ids: list[str] | None = None
    mode: str = "replace"  # replace | append


class KnowledgeItemCleanupRequest(BaseModel):
    pipeline: str = "流水线1"
    dry_run: bool = False
    delete_unknown_cause: bool = True
    delete_noise: bool = True


class KnowledgeItemAutoFillRequest(BaseModel):
    pipeline: str = "流水线1"
    limit: int = 60
    dry_run: bool = False


class PipelineCreateRequest(BaseModel):
    pipeline: str


def _extract_terms_from_text(text: str, top_k: int = 5) -> list[str]:
    # 仅抽取中文术语，避免知识图谱出现英文节点
    tokens = re.findall(r"[\u4e00-\u9fff]{2,}", text or "")
    stop = {
        "系统", "设备", "故障", "可以", "以及", "进行", "用于", "这个", "那个",
        "相关", "通过", "对于", "其中", "需要", "包括", "检查", "操作", "手册"
    }
    cleaned = [t for t in tokens if t not in stop]
    if not cleaned:
        return []
    return [w for w, _ in Counter(cleaned).most_common(top_k)]


def _short_text(s: str, max_len: int = 22) -> str:
    s = re.sub(r"\s+", "", s or "")
    return s if len(s) <= max_len else s[:max_len] + "…"


def _norm_text(s: str) -> str:
    return re.sub(r"[\s，。,.、；;：:【】\[\]()（）\-—_]+", "", (s or "").strip().lower())


def _machine_dedupe_key(name: str) -> str:
    raw = str(name or "").strip()
    if not raw:
        return ""
    compact = re.sub(r"[^0-9a-zA-Z\u4e00-\u9fff]+", "", raw)
    upper = compact.upper()
    codes = re.findall(r"\d+[A-Z]{1,4}\d+", upper)
    if codes:
        return f"model:{codes[0]}"
    return _norm_text(raw) or raw.lower()


def _normalize_machine_name(name: str) -> str:
    raw = str(name or "").strip()
    if not raw:
        return ""
    compact = re.sub(r"\s+", "", raw).upper()
    if "1FT7" in compact:
        return "SIMOTICS S-1FT7"
    return raw


def _pick_preferred_machine_name(names) -> str:
    arr = [str(x or "").strip() for x in (names or [])]
    arr = [x for x in arr if x]
    if not arr:
        return "设备"

    normalized = [_normalize_machine_name(x) for x in arr]
    if any(x == "SIMOTICS S-1FT7" for x in normalized):
        return "SIMOTICS S-1FT7"

    def score(v: str) -> tuple[int, int, str]:
        lo = v.lower()
        s = 0
        if "simotics" in lo:
            s += 50
        if "siemens" in lo:
            s += 50
        if "-" in v or "_" in v:
            s += 5
        if " " in v:
            s += 2
        if "系列" in v:
            s -= 2
        return (s, len(v), v)

    return max(arr, key=score)


def _is_noise_phrase(s: str) -> bool:
    t = _norm_text(s)
    if not t:
        return True
    if t.isdigit():
        return True
    if "手册" in t or "目录" in t:
        return True
    if any(_norm_text(x) in t for x in NOISE_TERMS):
        return True
    return False


def _clean_fault_name(s: str) -> str:
    v = _short_text(s, 26)
    if _is_noise_phrase(v):
        return ""
    return v


def _clean_solution(s: str) -> str:
    v = _short_text(s, 36)
    if _is_noise_phrase(v):
        return ""
    if not any(k in v for k in SOLUTION_HINTS):
        return ""
    return v


def _canonicalize_device_name(name: str) -> str:
    raw = str(name or "").strip()
    if not raw:
        return ""
    compact = re.sub(r"\s+", "", raw)
    for canonical, aliases in CANONICAL_DEVICE_ALIASES:
        if any(alias in compact for alias in aliases):
            return canonical
    return _short_text(compact, 24)


def _is_generic_fault_phrase(s: str) -> bool:
    raw = str(s or "").strip()
    t = _norm_text(raw)
    if not t:
        return True
    generic_exact = {
        "故障", "设备故障", "设备异常", "系统故障", "系统异常",
        "电饭煲故障", "吸尘器故障", "传送带故障", "输送带故障",
        "电饭煲异常", "吸尘器异常", "传送带异常", "输送带异常",
    }
    if t in {_norm_text(x) for x in generic_exact}:
        return True
    if re.fullmatch(r"故障[0-9一二三四五六七八九十]+", raw):
        return True
    if re.fullmatch(r"[0-9一二三四五六七八九十]+", raw):
        return True
    canonical_devices = {c for c, _aliases in CANONICAL_DEVICE_ALIASES}
    for dev in canonical_devices:
        dn = _norm_text(dev)
        if t in {dn, dn + "故障", dn + "异常"}:
            return True
    return False


def _is_valid_fault_phrase(s: str) -> bool:
    raw = str(s or "").strip()
    t = _norm_text(raw)
    if not t:
        return False
    if _is_generic_fault_phrase(raw):
        return False
    if len(raw) < 2 or len(raw) > 30:
        return False
    if any(x in raw for x in ["手册", "目录", "图示", "步骤", "注意事项", "热线", "更新日期"]):
        return False
    if any(x in raw for x in ["若", "如果", "则", "需要", "应当", "建议", "请", "确认"]):
        return False
    if any(k in raw for k in SOLUTION_HINTS):
        return False
    if "，" in raw or "," in raw or "。" in raw:
        return False
    if any(x in t for x in ["排查表", "技术参数", "保养计划", "维修指南", "查询热线", "更新日期"]):
        return False
    if any(x in t for x in ["每日", "每周", "每月", "每半年", "每年"]):
        return False
    allowed_short = set(COMMON_FAULT_TITLES)
    symptom_patterns = [
        r"无法|不能|不通电|无反应|不加热|不熟|煮糊|溢出|失灵|错误代码|吸力减弱|异常噪音|充电故障|无法开机|不启动|无法充电|充不进电"
    ]
    if any(x in s for x in FAULT_HINTS) or s in allowed_short:
        return True
    if any(re.search(p, raw) for p in symptom_patterns):
        return True
    return False


def _extract_graph_from_content(content: str, main_device: str | None = None) -> dict:
    sentences = [x.strip() for x in re.split(r"[。！？；;\n]+", content or "") if x.strip()]
    graph = {}
    for sent in sentences:
        has_fault = any(k in sent for k in FAULT_HINTS)
        has_solution = any(k in sent for k in SOLUTION_HINTS)
        devices = [d for d in DEVICE_TERMS if d in sent]
        if main_device and (has_fault or has_solution):
            devices = [main_device]
        if not devices:
            continue
        fault_name = _short_text(sent)
        for dev in devices:
            graph.setdefault(dev, {"faults": {}})
            if has_fault:
                graph[dev]["faults"].setdefault(fault_name, {"solutions": []})
            if has_solution:
                target_fault = fault_name if has_fault else (next(iter(graph[dev]["faults"]), None))
                if target_fault:
                    sol = _short_text(sent, 28)
                    if sol not in graph[dev]["faults"][target_fault]["solutions"]:
                        graph[dev]["faults"][target_fault]["solutions"].append(sol)
    return graph


def _extract_json_dict(text: str) -> dict:
    text = (text or "").strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if len(lines) >= 3 and lines[-1].strip().startswith("```"):
            text = "\n".join(lines[1:-1]).strip()
    m = re.search(r"\{[\s\S]*\}", text)
    if not m:
        raise ValueError("AI 输出中未找到 JSON")
    return json.loads(m.group())


def _sanitize_kg_payload(data: dict) -> dict:
    devices = data.get("devices", []) if isinstance(data, dict) else []
    cleaned_devices = []
    for dev in devices:
        name = _canonicalize_device_name(str((dev or {}).get("name", "")).strip())
        if not name or _is_noise_phrase(name):
            continue
        faults = []
        fault_seen = set()
        for f in (dev or {}).get("faults", []) or []:
            fname = _clean_fault_name(str((f or {}).get("name", "")).strip())
            if not fname:
                continue
            if not _is_valid_fault_phrase(fname):
                continue
            sols = []
            sol_seen = set()
            for s in (f or {}).get("solutions", []) or []:
                sv = _clean_solution(str(s).strip())
                nk = _norm_text(sv)
                if sv and nk not in sol_seen:
                    sols.append(sv)
                    sol_seen.add(nk)
            if not sols:
                continue
            fk = _norm_text(fname)
            if fk in fault_seen:
                continue
            fault_seen.add(fk)
            if any(_norm_text(x) in fk for x in NOISE_TERMS):
                continue
            faults.append({"name": fname, "solutions": sols[:6]})
        if faults:
            cleaned_devices.append({"name": name, "faults": faults[:10]})
    return {"devices": cleaned_devices[:30]}


def _has_useful_kg(kg: dict) -> bool:
    devices = (kg or {}).get("devices", []) if isinstance(kg, dict) else []
    for d in devices:
        for f in (d or {}).get("faults", []) or []:
            sols = (f or {}).get("solutions", []) or []
            if str((f or {}).get("name", "")).strip() and len(sols) > 0:
                return True
    return False


def _infer_main_device(content: str) -> str:
    text = content or ""
    m = re.search(r"([\u4e00-\u9fff]{2,20})(?:维修保养手册|维修手册|保养手册)", text)
    if m:
        return _canonicalize_device_name(m.group(1))
    for d in sorted(DEVICE_TERMS, key=len, reverse=True):
        if d in text:
            return _canonicalize_device_name(d)
    return "设备"


def _infer_device_from_filename(name: str) -> str:
    raw = str(name or "").strip()
    if not raw:
        return "设备"
    n = re.sub(r"\.[^.]+$", "", raw)
    m = re.search(r"([\u4e00-\u9fff]{2,40})(维修保养手册|维修手册|保养手册|操作说明|说明书)?", n)
    return (m.group(1) if m else (n or "设备")).strip() or "设备"


def _infer_machine_category_from_machine(machine: str) -> str:
    v = re.sub(r"\s+", "", str(machine or "")).lower()
    if not v:
        return ""
    return _canonicalize_machine_category("", v)


def _canonicalize_machine_category(category: str | None, machine: str | None = None) -> str:
    raw_input = str(category or "").strip()
    raw = re.sub(r"\s+", "", raw_input).lower()
    machine_text = re.sub(r"\s+", "", str(machine or "")).lower()
    text = raw or machine_text
    if not text:
        return ""

    motor_keys = [
        "伺服电机", "同步电机", "异步电机", "步进电机", "直流电机", "交流电机",
        "减速电机", "力矩电机", "主轴电机", "马达", "motor", "servo", "同步机", "异步机"
    ]
    if any(k in text for k in motor_keys) or ("电机" in text):
        return "电机"

    if not raw:
        for canonical, aliases in CANONICAL_DEVICE_ALIASES:
            if canonical and canonical in machine_text:
                return canonical
            for a in (aliases or []):
                if a and (re.sub(r"\s+", "", str(a)).lower() in machine_text):
                    return canonical

    mapping = [
        ("变频器", ["变频", "inverter"]),
        ("PLC", ["plc"]),
        ("传送带", ["传送带", "输送带"]),
        ("传感器", ["传感器", "sensor"]),
        ("轴承", ["轴承", "bearing"]),
        ("液压", ["液压", "hydraulic"]),
        ("气动", ["气动", "pneumatic"]),
    ]
    for cat, keys in mapping:
        if any(k in text for k in keys):
            return cat
    if raw:
        return _clean_single_line_text(raw_input, 120)
    return ""


def _infer_problem_category_from_problem(problem: str) -> str:
    v = re.sub(r"\s+", "", str(problem or "")).lower()
    if not v:
        return ""
    return _canonicalize_problem_category("", problem, "")


def _canonicalize_problem_category(category: str | None, problem: str | None = None, root_cause: str | None = None) -> str:
    raw_input = str(category or "").strip()
    raw = _norm_text(raw_input)
    p = _norm_text(problem or "")
    rc = _norm_text(root_cause or "")
    text = raw or (p + rc)
    if not text:
        return ""

    if any(k in text for k in ["短路", "断路", "漏电", "跳闸", "过流", "过压", "欠压", "电源", "接线", "电池", "充电", "电路", "电机", "烧毁", "保险丝"]):
        return "电气"
    if any(k in text for k in ["报警", "通讯", "通信", "程序", "参数", "plc", "伺服", "驱动器", "编码器", "控制器", "协议"]):
        return "控制"
    if any(k in text for k in ["液压", "油", "泄漏", "压力", "泵", "阀", "油温"]):
        return "液压"
    if any(k in text for k in ["气动", "气压", "气缸", "电磁阀", "漏气"]):
        return "气动"
    if any(k in text for k in ["传感器", "信号", "误报警", "漂移", "采样", "探头"]):
        return "传感器"
    if any(k in text for k in ["温度", "湿度", "粉尘", "盐雾", "腐蚀", "环境", "高温", "低温", "振动环境"]):
        return "环境"
    if any(k in text for k in ["保养", "维护", "清洁", "清理", "更换", "润滑", "检修", "校准", "清洗"]):
        return "维护"
    if any(k in text for k in ["使用", "误操作", "操作", "未关机", "不规范", "违规", "安装错误", "接错", "未按"]):
        return "使用"
    if any(k in text for k in ["配件", "非原装", "不匹配", "型号不对", "规格不对", "兼容", "替代件"]):
        return "配件"
    if any(k in text for k in ["吸力", "效率", "性能", "功率下降", "无力", "速度慢"]):
        return "性能"
    if any(k in text for k in ["振动", "异响", "磨损", "断裂", "卡滞", "堵塞", "松动", "轴承", "摩擦", "碰撞", "变形", "裂纹", "刮擦"]):
        return "机械"
    if raw:
        return _clean_single_line_text(raw_input, 120)
    return "其他"


def _pick_reliable_main_device(base: dict, content: str, title_hint: str = "") -> str:
    inferred = _infer_main_device((title_hint or "") + "\n" + (content or "")[:2000])
    if inferred and inferred != "设备":
        return inferred
    text = (content or "")[:6000]
    vacuum_score = sum(1 for kw in ["吸力", "集尘", "滤网", "刷头", "吸管", "充电器", "电池电量", "风道"] if kw in text)
    rice_score = sum(1 for kw in ["电饭煲", "加热盘", "内锅", "煮饭", "磁钢", "限温器", "内胆", "上盖"] if kw in text)
    if vacuum_score >= 2 and vacuum_score > rice_score:
        return "吸尘器"
    if rice_score >= 2 and rice_score > vacuum_score:
        return "电饭煲"
    best_name = ""
    best_fault_count = 0
    for d in (base or {}).get("devices", []) or []:
        name = _canonicalize_device_name(str((d or {}).get("name", "")).strip())
        if not name or name == "设备":
            continue
        fault_count = len((d or {}).get("faults", []) or [])
        if fault_count > best_fault_count:
            best_name = name
            best_fault_count = fault_count
    return best_name


def _extract_numbered_fault_blocks(content: str, limit: int = 12) -> list[dict]:
    text = content or ""
    pattern = re.compile(r"故障\s*[0-9一二三四五六七八九十]+\s*[：:]\s*([^\n]{2,48})")
    matches = list(pattern.finditer(text))
    if not matches:
        return []
    pairs = []
    for i, m in enumerate(matches):
        fault_raw = m.group(1).strip()
        fault_raw = re.split(r"[（(]", fault_raw)[0].strip() or fault_raw
        fault_name = _clean_fault_name(fault_raw)
        if not fault_name or not _is_valid_fault_phrase(fault_name):
            continue
        st = m.end()
        ed = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        block = text[st:ed]
        step_match = re.search(r"(维修步骤|解决方法|处理步骤)\s*[：:]([\s\S]*)", block)
        seg = step_match.group(2) if step_match else block
        lines = []
        for ln in re.split(r"[\n\r]+", seg):
            for part in re.split(r"[；;。]", ln):
                p = re.sub(r"^[\s·\-•\d\.\、]+\s*", "", part.strip())
                if p:
                    lines.append(p)
        solutions = []
        seen = set()
        for ln in lines:
            sv = _clean_solution(ln)
            if not sv:
                continue
            nk = _norm_text(sv)
            if nk in seen:
                continue
            seen.add(nk)
            solutions.append(sv)
            if len(solutions) >= 6:
                break
        if solutions:
            pairs.append({"name": fault_name, "solutions": solutions})
        if len(pairs) >= limit:
            break
    return pairs


def _constrain_kg_to_main_device(kg: dict, content: str, title_hint: str = "") -> dict:
    base = _sanitize_kg_payload(kg or {})
    if not isinstance(base, dict):
        return {"devices": []}
    main_dev = _pick_reliable_main_device(base, content, title_hint)
    if not main_dev:
        return base
    all_faults = []
    seen_fault = set()
    for d in base.get("devices", []) or []:
        for f in (d or {}).get("faults", []) or []:
            fname = _clean_fault_name(str((f or {}).get("name", "")).strip())
            if not fname or not _is_valid_fault_phrase(fname):
                continue
            fk = _norm_text(fname)
            if fk in seen_fault:
                continue
            sols = []
            seen_sol = set()
            for s in (f or {}).get("solutions", []) or []:
                sv = _clean_solution(str(s).strip())
                nk = _norm_text(sv)
                if sv and nk not in seen_sol:
                    sols.append(sv)
                    seen_sol.add(nk)
            if not sols:
                continue
            all_faults.append({"name": fname, "solutions": sols[:6]})
            seen_fault.add(fk)
    if not all_faults:
        return {"devices": []}
    return {"devices": [{"name": main_dev, "faults": all_faults[:12]}]}


def _extract_fault_solution_pairs_from_content(content: str, limit: int = 10) -> list[dict]:
    lines = [re.sub(r"^[\s·\-•\d\.\、]+", "", x.strip()) for x in (content or "").splitlines()]
    lines = [x for x in lines if x]
    pairs = []
    i = 0
    while i < len(lines):
        line = lines[i]
        is_fault = (2 <= len(line) <= 16) and (
            any(k in line for k in FAULT_HINTS) or line in {"无法开机", "吸力减弱", "异常噪音", "充电故障"}
        )
        if not is_fault:
            i += 1
            continue
        fault = _short_text(line, 24)
        sols = []
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            next_is_fault = (2 <= len(nxt) <= 16) and (
                any(k in nxt for k in FAULT_HINTS) or nxt in {"无法开机", "吸力减弱", "异常噪音", "充电故障"}
            )
            if next_is_fault:
                break
            if any(k in nxt for k in SOLUTION_HINTS):
                sols.append(_short_text(nxt, 32))
            j += 1
        if sols:
            pairs.append({"name": fault, "solutions": list(dict.fromkeys(sols))[:6]})
        i = j
        if len(pairs) >= limit:
            break
    return pairs


def _extract_fault_solution_by_position(content: str, limit: int = 12) -> list[dict]:
    text = re.sub(r"\s+", " ", content or "")
    if not text:
        return []

    exact_hits = []
    for t in COMMON_FAULT_TITLES:
        for m in re.finditer(re.escape(t), text):
            exact_hits.append((m.start(), m.end(), t))

    hits = []
    if len(exact_hits) >= 2:
        hits = exact_hits
    else:
        hits.extend(exact_hits)
        for m in re.finditer(r"[\u4e00-\u9fff]{2,12}(?:故障|异常|报警|失效|过热|异响|堵塞|磨损|卡滞|短路|断路|跳闸|停机)", text):
            hits.append((m.start(), m.end(), m.group(0)))
    if not hits:
        return []

    hits.sort(key=lambda x: x[0])
    merged = []
    seen = set()
    for st, ed, name in hits:
        k = (st // 2, _norm_text(name))
        if k in seen:
            continue
        seen.add(k)
        merged.append((st, ed, _clean_fault_name(name)))
    merged = [(a, b, c) for (a, b, c) in merged if c and _is_valid_fault_phrase(c)]
    if not merged:
        return []

    pairs = []
    for i, (st, _ed, fault) in enumerate(merged):
        right = merged[i + 1][0] if i + 1 < len(merged) else len(text)
        seg = text[st:right]
        sols = []
        for m in re.finditer(rf"[^。；\n]{{0,18}}(?:{'|'.join(SOLUTION_HINTS)})[^。；\n]{{0,30}}", seg):
            sv = _clean_solution(m.group(0))
            if sv:
                nk = _norm_text(sv)
                if nk not in {_norm_text(x) for x in sols}:
                    sols.append(sv)
        if not sols:
            for m in re.finditer(r"\d+[\.、]\s*([^。；\n]{4,40})", seg):
                sv = _clean_solution(m.group(1))
                if sv:
                    nk = _norm_text(sv)
                    if nk not in {_norm_text(x) for x in sols}:
                        sols.append(sv)
        if sols:
            pairs.append({"name": fault, "solutions": sols[:6]})
        if len(pairs) >= limit:
            break
    return pairs

def _parse_table_kg(content: str, limit_faults: int = 12) -> dict:
    text = (content or "").replace("\t", " ").replace("｜", "|")
    pattern = re.compile(r"(?P<f>[^|\n]{2,30})\s*\|\s*(?P<r>[^|\n]{1,200})\s*\|\s*(?P<s>[^|\n]{2,240})")
    rows = []
    for m in pattern.finditer(text):
        f = m.group("f").strip()
        r = m.group("r").strip()
        s = m.group("s").strip()
        h = (f + r + s).replace(" ", "")
        if "故障现象" in h and "解决方法" in h:
            continue
        if not f or not s:
            continue
        rows.append((f, s))
        if len(rows) >= limit_faults * 2:
            break
    if not rows:
        # 回退到逐行方式
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        for ln in lines:
            segs = [s.strip() for s in re.split(r"\s*\|\s*", ln) if s.strip()]
            if len(segs) >= 3:
                f, s = segs[0], segs[-1]
                h = (segs[0] + "".join(segs[1:])).replace(" ", "")
                if "故障现象" in h and "解决方法" in h:
                    continue
                rows.append((f, s))
            if len(rows) >= limit_faults * 2:
                break
        if not rows:
            return {"devices": []}
    def clean_num_prefix(s: str) -> str:
        return re.sub(r"^\d+[\.、]\s*", "", s.strip())
    pairs = []
    seen_fault = set()
    for f, s in rows:
        f1 = clean_num_prefix(f)
        f1 = _clean_fault_name(f1)
        if not f1 or not _is_valid_fault_phrase(f1):
            continue
        fk = _norm_text(f1)
        if fk in seen_fault:
            continue
        sols = []
        for part in re.split(r"[；;。]|(?<=\))\s*", s):
            for seg in re.split(r"\d+[\.、]\s*", part):
                sv = _clean_solution(seg)
                if sv:
                    nk = _norm_text(sv)
                    if nk not in {_norm_text(x) for x in sols}:
                        sols.append(sv)
        if not sols:
            continue
        pairs.append({"name": f1, "solutions": sols[:6]})
        seen_fault.add(fk)
        if len(pairs) >= limit_faults:
            break
    if not pairs:
        return {"devices": []}
    dev = _infer_main_device(content)
    return {"devices": [{"name": dev, "faults": pairs}]}


def _extract_table_pairs_from_content(content: str, limit: int = 12) -> list[dict]:
    raw = content or ""
    start = 0
    m_start = re.search(r"(常见故障排查表|故障排查表)", raw)
    if m_start:
        start = m_start.start()
    end = len(raw)
    m_end = re.search(r"(分步维修指南|定期保养计划|技术参数|注意事项)", raw[start:])
    if m_end:
        end = start + m_end.start()
    focus = raw[start:end]

    lines = [re.sub(r"^[\s·\-•\d\.\、]+", "", x.strip()) for x in focus.splitlines()]
    lines = [x for x in lines if x]
    pairs = []
    i = 0
    while i < len(lines):
        line = lines[i]
        fault = _clean_fault_name(line)
        is_fault = bool(fault) and (2 <= len(fault) <= 16) and (
            any(k in fault for k in FAULT_HINTS) or fault in {"无法开机", "吸力减弱", "异常噪音", "充电故障"}
        )
        if is_fault and not _is_valid_fault_phrase(fault):
            is_fault = False
        if not is_fault:
            i += 1
            continue
        j = i + 1
        solutions = []
        while j < len(lines):
            nxt = lines[j]
            nxt_fault = _clean_fault_name(nxt)
            next_is_fault = bool(nxt_fault) and (2 <= len(nxt_fault) <= 16) and (
                any(k in nxt_fault for k in FAULT_HINTS) or nxt_fault in {"无法开机", "吸力减弱", "异常噪音", "充电故障"}
            )
            if next_is_fault:
                break
            sv = _clean_solution(nxt)
            if sv:
                key = _norm_text(sv)
                if key not in {_norm_text(x) for x in solutions}:
                    solutions.append(sv)
            j += 1
        if solutions:
            pairs.append({"name": fault, "solutions": solutions[:6]})
        i = j
        if len(pairs) >= limit:
            break
    return pairs


def _build_fallback_kg(content: str) -> dict:
    def _extract_known_faults_pairs(text: str) -> list[dict]:
        t = text or ""
        idxs = []
        for name in ["无法开机", "吸力减弱", "异常噪音", "充电故障"]:
            for m in re.finditer(re.escape(name), t):
                idxs.append((m.start(), name))
        if not idxs:
            return []
        idxs.sort()
        pairs = []
        for i, (st, name) in enumerate(idxs):
            ed = idxs[i + 1][0] if i + 1 < len(idxs) else len(t)
            seg = t[st:ed]
            sols = []
            for ln in re.split(r"[\n；;。]", seg):
                ln = re.sub(r"^[\s·\-•\d\.\、]+\s*", "", ln.strip())
                sv = _clean_solution(ln)
                if sv:
                    nk = _norm_text(sv)
                    if nk not in {_norm_text(x) for x in sols}:
                        sols.append(sv)
                if len(sols) >= 5:
                    break
            if sols:
                pairs.append({"name": name, "solutions": sols})
        return pairs

    text = content[:120000]
    main_dev = _infer_main_device(content)
    candidates = []

    pf = _extract_known_faults_pairs(text)
    if pf:
        candidates.append({"devices": [{"name": main_dev, "faults": pf}]})
    numbered_pairs = _extract_numbered_fault_blocks(text)
    if numbered_pairs:
        candidates.append({"devices": [{"name": main_dev, "faults": numbered_pairs}]})

    table_kg = _parse_table_kg(text)
    if _has_useful_kg(table_kg):
        candidates.append(_constrain_kg_to_main_device(table_kg, text))
    graph = _extract_graph_from_content(content[:120000], main_device=main_dev)
    devices = []
    for dev, val in graph.items():
        faults = []
        for fault, fval in val["faults"].items():
            if fval["solutions"]:
                faults.append({"name": fault, "solutions": fval["solutions"][:5]})
        if faults:
            devices.append({"name": dev, "faults": faults[:8]})

    pairs = _extract_fault_solution_by_position(text)
    if not pairs:
        pairs = _extract_numbered_fault_blocks(text)
    if not pairs:
        pairs = _extract_table_pairs_from_content(text)
    if not pairs:
        pairs = _extract_fault_solution_pairs_from_content(text)
    if pairs:
        found = next((d for d in devices if d["name"] == main_dev), None)
        if found:
            exists = {x["name"] for x in found["faults"]}
            for p in pairs:
                if p["name"] not in exists:
                    found["faults"].append(p)
        else:
            devices.insert(0, {"name": main_dev, "faults": pairs[:10]})
    if devices:
        candidates.append({"devices": devices[:20]})

    if not candidates:
        return {"devices": []}

    def score(kg: dict) -> tuple[int, int]:
        ds = kg.get("devices", []) if isinstance(kg, dict) else []
        fault_cnt = sum(len((d or {}).get("faults", []) or []) for d in ds)
        sol_cnt = sum(len((f or {}).get("solutions", []) or []) for d in ds for f in (d or {}).get("faults", []) or [])
        return fault_cnt, sol_cnt

    best = max(candidates, key=score)
    return _constrain_kg_to_main_device(best, text)


def _build_minimum_kg(content: str) -> dict:
    text = content or ""
    compact = re.sub(r"\s+", "", text)
    main_dev = _infer_main_device(compact or text)
    fault_patterns = [
        r"[\u4e00-\u9fff]{1,10}(?:无法开机|吸力减弱|异常噪音|充电故障|故障|异常|过热|异响|堵塞|磨损|卡滞|短路|断路|跳闸|停机|无压力)"
    ]
    faults = []
    seen = set()
    for p in fault_patterns:
        for m in re.findall(p, compact or text):
            name = _clean_fault_name(m)
            if not name or not _is_valid_fault_phrase(name):
                continue
            k = _norm_text(name)
            if k in seen:
                continue
            seen.add(k)
            faults.append(name)
            if len(faults) >= 8:
                break
        if len(faults) >= 8:
            break

    solution_lines = []
    for line in ((text.splitlines() if "\n" in text else re.split(r"[。；;\n]+", text))):
        line = re.sub(r"^[\s·\-•\d\.\、]+", "", line.strip())
        sv = _clean_solution(line)
        if sv:
            nk = _norm_text(sv)
            if nk not in {_norm_text(x) for x in solution_lines}:
                solution_lines.append(sv)
        if len(solution_lines) >= 20:
            break

    if not faults and solution_lines:
        faults = ["设备运行异常"]
    if not faults:
        return {"devices": []}

    pairs = []
    for i, f in enumerate(faults):
        sols = solution_lines[i * 2:(i + 1) * 2] or solution_lines[:2]
        if not sols:
            continue
        pairs.append({"name": f, "solutions": sols[:5]})
    if not pairs:
        return {"devices": []}
    return {"devices": [{"name": main_dev, "faults": pairs[:8]}]}


def _merge_pairs_into_kg(kg: dict, content: str) -> dict:
    base = _sanitize_kg_payload(kg or {})
    pairs = _extract_table_pairs_from_content(content)
    if not pairs:
        return base
    main_dev = _infer_main_device(content)
    devices = list(base.get("devices", []))
    target = next((d for d in devices if _norm_text(d.get("name", "")) == _norm_text(main_dev)), None)
    if not target:
        target = {"name": main_dev, "faults": []}
        devices.insert(0, target)
    exists = {_norm_text(x.get("name", "")) for x in target.get("faults", [])}
    for p in pairs:
        nk = _norm_text(p.get("name", ""))
        if nk and nk not in exists:
            target["faults"].append(p)
            exists.add(nk)
    base["devices"] = devices[:30]
    return _constrain_kg_to_main_device(base, content)


async def _build_doc_kg_with_ai(content: str) -> dict:
    text = (content or "")[:120000]
    if not text:
        return {"devices": []}
    prompt = f"""
你是工业设备维修知识图谱抽取器，负责从维修手册中提取“设备 → 故障问题 → 解决方案”。

任务重点
- 优先识别“常见故障排查表/故障现象/可能原因/解决方法”三列表格
- 容错粘贴/OCR 格式：列可能错位、换行、使用数字序号（1. 2. 3.）
- 输出只保留“有故障且有明确解决动作”的条目；忽略噪声性标题或周期性维护项

强制约束
1. 仅输出 JSON；不要输出任何解释文字
2. 全部中文
3. 设备名称：若标题类似“XX维修保养手册/维修手册/保养手册”，设备名用“XX”
4. 故障优先从表格“故障现象”列逐行抽取；常见模板示例：
   故障现象 | 可能原因 | 解决方法
   无法开机 | ...     | ...
   吸力减弱 | ...     | ...
   异常噪音 | ...     | ...
   充电故障 | ...     | ...
5. 解决方案必须是“动作句”，如：检查/更换/清理/调整/润滑/测试/校准/复位/重启/排查
6. 去重与清洗：
   - 丢弃“常见故障排查表/技术参数/定期保养计划/安全须知/产品结构图示/分步维修指南/每日/每周/每月/每半年/每年/维修查询热线/更新日期/可能原因/解决方法/故障现象”等标题或周期词
   - 丢弃无动作动词的方案
   - 严禁把“原因句/条件句/判断句”当作故障名，例如包含“若/如果/则/需/需要/确认/建议”的句子不能作为故障名
   - 严禁把“手册标题/章节标题/目录项”当作故障名
7. 故障字段要求：
   - 必须是“故障现象短语”，长度 2~16 字，不能含逗号句号
   - 优先使用：无法开机/吸力减弱/异常噪音/充电故障/无法启动/不启动/无法充电/充不进电/过热报警/异响/无压力/压力不足
   - 若出现“温度传感器故障/加热盘损坏”等也可保留
8. 解决方案字段要求：
   - 必须是可执行动作，不得只是“原因判断”
   - 如果同一行只有“原因”没有动作，则不输出该方案
9. 数量要求：
   - 每个设备输出 3~8 个故障（如果手册中确实只有更少，则按实际）
   - 每个故障输出 1~5 条解决方案
10. 若未识别到表格：
   - 在全文中查找明确的故障标题（例如：无法开机/吸力减弱/异常噪音/充电故障/无法启动/不启动/无法充电/充不进电/过热报警/异响/无压力/压力不足），
     并在相邻段落或该标题下方提取动作型解决方案
11. 如仍得不到有效条目，返回空数组 devices: []

输出结构（严格遵循）：
{{
  "devices": [
    {{
      "name": "设备名",
      "faults": [
        {{
          "name": "故障现象",
          "solutions": ["解决方案1", "解决方案2", "解决方案3"]
        }}
      ]
    }}
  ]
}}

维修手册内容：
{text}
"""
    manager = get_llm_manager()
    resp, _provider = await manager.generate_with_fallback(prompt)
    parsed = _extract_json_dict(resp.content if resp else "")
    return _sanitize_kg_payload(parsed)


@router.post("/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    pipeline: str = Form("流水线1"),
    auto_extract: bool = Form(True),
    background_tasks: BackgroundTasks = None,
    user: dict = Depends(require_expert),
):
    """上传设备文档，自动解析分块并存入 PostgreSQL 向量库"""
    # 文件大小校验
    if file.size and file.size > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小超过50MB")

    allowed = {".pdf", ".txt", ".log", ".docx"}
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")
    pipeline = _normalize_pipeline(pipeline)

    # 生成文档ID
    doc_id = uuid.uuid4()

    # 保存原始文件到本地
    save_path = Path(settings.MANUALS_PATH) / f"{doc_id}_{file.filename}"
    save_path.parent.mkdir(parents=True, exist_ok=True)

    content = await file.read()
    async with aiofiles.open(save_path, "wb") as f:
        await f.write(content)

    # 解析文档
    try:
        chunks = parse_document(str(save_path))
    except Exception as e:
        save_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"文档解析失败: {e}")

    if not chunks:
        save_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="文档中未提取到有效文本")

    # 创建文档记录（psycopg2 直连）
    import psycopg2.extras
    try:
        with psycopg2.connect(
            host=settings.DB_HOST, port=settings.DB_PORT,
            user=settings.DB_USER, password=settings.DB_PASSWORD,
            database=settings.DB_NAME
        ) as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    INSERT INTO documents (doc_id, filename, file_size, file_type, status, metadata)
                    VALUES (%s, %s, %s, %s, 'processing', %s)
                """, (
                    str(doc_id), file.filename, len(content),
                    ext[1:], psycopg2.extras.Json({"original_path": str(save_path), "pipeline": pipeline})
                ))
                conn.commit()
    except Exception as e:
        save_path.unlink(missing_ok=True)
        raise HTTPException(status_code=503, detail=f"数据库不可用或写入失败: {e}")

    # 估算 token 数（用于成本记录）
    try:
        enc = tiktoken.get_encoding("cl100k_base")
        for chunk in chunks:
            chunk["tokens"] = len(enc.encode(chunk["text"]))
    except Exception:
        for chunk in chunks:
            chunk["tokens"] = None

    # 存入 PostgreSQL（文本 + 向量双写）
    try:
        merged_text = "\n".join((c.get("text", "") for c in chunks))
        try:
            kg_payload = await _build_doc_kg_with_ai(merged_text)
            kg_payload = _merge_pairs_into_kg(kg_payload, merged_text)
            if not _has_useful_kg(kg_payload):
                kg_payload = _build_fallback_kg(merged_text)
            if not _has_useful_kg(kg_payload):
                kg_payload = _build_minimum_kg(merged_text)
            kg_payload = _constrain_kg_to_main_device(kg_payload, merged_text, file.filename or "")
        except Exception:
            kg_payload = _build_fallback_kg(merged_text)
            if not _has_useful_kg(kg_payload):
                kg_payload = _build_minimum_kg(merged_text)
            kg_payload = _constrain_kg_to_main_device(kg_payload, merged_text, file.filename or "")

        res = await add_chunks_to_db(chunks, str(doc_id))
        embedded = bool(res.get("embedded"))
        with psycopg2.connect(
            host=settings.DB_HOST, port=settings.DB_PORT,
            user=settings.DB_USER, password=settings.DB_PASSWORD,
            database=settings.DB_NAME
        ) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    UPDATE documents
                    SET status = %s,
                        metadata = COALESCE(metadata, '{}'::jsonb) || %s::jsonb
                    WHERE doc_id = %s
                    """,
                    (
                        'active',
                        psycopg2.extras.Json({
                            "embedding": "ok" if embedded else "skipped",
                            "kg": kg_payload
                        }),
                        str(doc_id),
                    )
                )
                conn.commit()

        if auto_extract:
            async def _structured_extract_task(_doc_id: str, _pipeline: str):
                import psycopg2.extras
                if (settings.LLM_PROVIDER or "").lower() == "minimax" and (not settings.MINIMAX_API_KEY or not settings.MINIMAX_GROUP_ID):
                    with psycopg2.connect(
                        host=settings.DB_HOST, port=settings.DB_PORT,
                        user=settings.DB_USER, password=settings.DB_PASSWORD,
                        database=settings.DB_NAME
                    ) as conn2:
                        with conn2.cursor() as cur2:
                            cur2.execute(
                                """
                                UPDATE documents
                                SET metadata = COALESCE(metadata, '{}'::jsonb) || %s::jsonb
                                WHERE doc_id = %s
                                """,
                                (psycopg2.extras.Json({"structured_kb": "failed", "structured_error": "LLM 未配置"}), _doc_id),
                            )
                            conn2.commit()
                    return

                try:
                    r = await extract_knowledge_items_with_ai(pipeline=_pipeline, doc_ids=[_doc_id])
                    payload = {
                        "structured_kb": "ok",
                        "structured_extracted": int(getattr(r, "extracted", 0)),
                        "structured_inserted": int(getattr(r, "inserted", 0)),
                        "structured_skipped": int(getattr(r, "skipped", 0)),
                    }
                except Exception as e:
                    payload = {"structured_kb": "failed", "structured_error": str(e)[:200]}

                try:
                    with psycopg2.connect(
                        host=settings.DB_HOST, port=settings.DB_PORT,
                        user=settings.DB_USER, password=settings.DB_PASSWORD,
                        database=settings.DB_NAME
                    ) as conn2:
                        with conn2.cursor() as cur2:
                            cur2.execute(
                                """
                                UPDATE documents
                                SET metadata = COALESCE(metadata, '{}'::jsonb) || %s::jsonb
                                WHERE doc_id = %s
                                """,
                                (psycopg2.extras.Json(payload), _doc_id),
                            )
                            conn2.commit()
                except Exception:
                    pass

            with psycopg2.connect(
                host=settings.DB_HOST, port=settings.DB_PORT,
                user=settings.DB_USER, password=settings.DB_PASSWORD,
                database=settings.DB_NAME
            ) as conn2:
                with conn2.cursor() as cur2:
                    cur2.execute(
                        """
                        UPDATE documents
                        SET metadata = COALESCE(metadata, '{}'::jsonb) || %s::jsonb
                        WHERE doc_id = %s
                        """,
                        (psycopg2.extras.Json({"structured_kb": "pending"}), str(doc_id)),
                    )
                    conn2.commit()

            if background_tasks is not None:
                background_tasks.add_task(_structured_extract_task, str(doc_id), pipeline)
            else:
                await _structured_extract_task(str(doc_id), pipeline)
    except ValueError as e:
        with psycopg2.connect(
            host=settings.DB_HOST, port=settings.DB_PORT,
            user=settings.DB_USER, password=settings.DB_PASSWORD,
            database=settings.DB_NAME
        ) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    UPDATE documents
                    SET status = 'failed',
                        metadata = COALESCE(metadata, '{}'::jsonb) || %s::jsonb
                    WHERE doc_id = %s
                    """,
                    (psycopg2.extras.Json({"error": str(e)}), str(doc_id))
                )
                conn.commit()
        raise HTTPException(status_code=400, detail=f"向量入库失败: {e}")
    except Exception as e:
        with psycopg2.connect(
            host=settings.DB_HOST, port=settings.DB_PORT,
            user=settings.DB_USER, password=settings.DB_PASSWORD,
            database=settings.DB_NAME
        ) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    UPDATE documents
                    SET status = 'failed',
                        metadata = COALESCE(metadata, '{}'::jsonb) || %s::jsonb
                    WHERE doc_id = %s
                    """,
                    (psycopg2.extras.Json({"error": str(e)}), str(doc_id))
                )
                conn.commit()
        raise HTTPException(status_code=500, detail=f"向量入库失败: {e}")

    return UploadResponse(
        doc_id=str(doc_id),
        filename=file.filename,
        chunk_count=len(chunks),
        status="success",
    )


@router.get("/list")
async def list_documents(user: dict = Depends(get_current_user)):
    """列出已上传的文档"""
    import psycopg2.extras
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor(name="list_docs") as cur:
            cur.execute("""
                SELECT doc_id, filename, file_size, file_type, upload_time, status, metadata
                FROM documents WHERE status <> 'deleted'
                ORDER BY upload_time DESC
            """)
            rows = cur.fetchall()

    return [
        {
            "doc_id": str(row[0]),
            "filename": row[1],
            "file_size": row[2],
            "file_type": row[3],
            "upload_time": row[4].isoformat() if row[4] else None,
            "status": row[5],
            "pipeline": (row[6] or {}).get("pipeline", "流水线1"),
            "structured_kb": (row[6] or {}).get("structured_kb", ""),
            "structured_extracted": (row[6] or {}).get("structured_extracted", 0),
            "structured_inserted": (row[6] or {}).get("structured_inserted", 0),
            "structured_skipped": (row[6] or {}).get("structured_skipped", 0),
            "structured_error": (row[6] or {}).get("structured_error", ""),
        }
        for row in rows
    ]


@router.delete("/{doc_id}")
async def delete_document(doc_id: str, user: dict = Depends(require_expert)):
    """软删除文档（将 status 改为 deleted，向量数据通过外键级联删除）"""
    import psycopg2.extras
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE documents SET status = 'deleted' WHERE doc_id = %s AND status <> 'deleted'",
                (doc_id,)
            )
            conn.commit()
            if cur.rowcount == 0:
                raise HTTPException(status_code=404, detail="文档不存在或已是删除状态")

    return {"message": "文档已删除", "doc_id": doc_id}


@router.put("/{doc_id}/pipeline")
async def update_document_pipeline(doc_id: str, pipeline: str, user: dict = Depends(require_expert)):
    pipeline = _normalize_pipeline(pipeline)

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                UPDATE documents
                SET metadata = COALESCE(metadata, '{}'::jsonb) || %s::jsonb
                WHERE doc_id = %s AND status <> 'deleted'
                """,
                (psycopg2.extras.Json({"pipeline": pipeline}), doc_id),
            )
            conn.commit()
            if cur.rowcount == 0:
                raise HTTPException(status_code=404, detail="文档不存在或已删除")

    return {"doc_id": doc_id, "pipeline": pipeline, "message": "流水线分组已更新"}


@router.get("/pipelines")
async def list_pipelines():
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        _ensure_pipeline_registry_table(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT DISTINCT pipeline FROM (
                    SELECT COALESCE(metadata->>'pipeline', '流水线1') AS pipeline
                    FROM documents
                    WHERE status <> 'deleted'
                    UNION ALL
                    SELECT pipeline
                    FROM knowledge_items
                    WHERE status <> 'deleted'
                    UNION ALL
                    SELECT pipeline
                    FROM pipeline_registry
                ) t
                WHERE pipeline IS NOT NULL AND pipeline <> ''
                ORDER BY pipeline
                """
            )
            rows = cur.fetchall()
    vals = [str(r[0]).strip() for r in rows if r and str(r[0]).strip()]
    if "流水线1" not in vals:
        vals.insert(0, "流水线1")
    return {"pipelines": vals}


@router.post("/pipelines")
async def create_pipeline(payload: PipelineCreateRequest, user: dict = Depends(require_expert)):
    pipeline = _normalize_pipeline(payload.pipeline)
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        _ensure_pipeline_registry_table(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO pipeline_registry(pipeline)
                VALUES (%s)
                ON CONFLICT (pipeline) DO NOTHING
                """,
                (pipeline,),
            )
            conn.commit()
    return {"pipeline": pipeline, "message": "流水线已创建"}


@router.get("/stats")
async def get_stats():
    """获取系统统计信息"""
    import psycopg2
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT COUNT(*) FROM documents WHERE status = 'active'")
            doc_count = cur.fetchone()[0]
            cur.execute("SELECT COUNT(*) FROM document_chunks")
            chunk_count = cur.fetchone()[0]
    return {
        "total_docs": doc_count or 0,
        "total_chunks": chunk_count or 0,
    }


@router.post("/search")
async def search_knowledge(query: str, top_k: int = 5):
    """在知识库中搜索，返回相关段落"""
    try:
        results = await retrieve(query, top_k=top_k)
        return {
            "results": results,
            "count": len(results),
            "query": query,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")


@router.post("/items")
async def create_knowledge_item(payload: KnowledgeItemCreateRequest, enrich: bool = False):
    pipeline = _normalize_pipeline(payload.pipeline)
    problem = _clean_single_line_text(payload.problem, 240)
    if not problem:
        raise HTTPException(status_code=400, detail="problem 不能为空")
    root_cause = _clean_single_line_text(payload.root_cause, 320)
    if not root_cause:
        raise HTTPException(status_code=400, detail="root_cause 不能为空")

    machine = _clean_single_line_text(payload.machine, 160)
    machine_category = _clean_single_line_text(payload.machine_category, 120)
    problem_category = _clean_single_line_text(payload.problem_category, 120)
    solution = _clean_single_line_text(payload.solution, 800)

    if enrich:
        if not machine_category:
            machine_category = _infer_machine_category_from_machine(machine) if machine else _infer_machine_category_from_machine(root_cause)
        if not problem_category:
            problem_category = _infer_problem_category_from_problem(problem)

        if (not machine_category) or (not problem_category):
            try:
                ai_out = await _ai_enrich_categories(pipeline, machine, problem, root_cause)
                if not machine_category:
                    machine_category = _canonicalize_machine_category(ai_out.get("machine_category", ""), machine)
                if not problem_category:
                    problem_category = _clean_single_line_text(ai_out.get("problem_category", ""), 120)
            except Exception:
                pass

    machine_category = _canonicalize_machine_category(machine_category, machine)
    problem_category = _canonicalize_problem_category(problem_category, problem, root_cause)
    _validate_knowledge_item_fields(pipeline, machine_category, machine, problem_category, problem, root_cause, solution)

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO knowledge_items (
                    pipeline, machine_category, machine, problem_category,
                    problem, root_cause, solution, metadata, status
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s::jsonb, 'active')
                ON CONFLICT (pipeline, machine, problem, root_cause)
                DO UPDATE SET
                    machine_category = EXCLUDED.machine_category,
                    problem_category = EXCLUDED.problem_category,
                    solution = CASE WHEN COALESCE(knowledge_items.solution, '') = '' THEN EXCLUDED.solution ELSE knowledge_items.solution END,
                    metadata = COALESCE(knowledge_items.metadata, '{}'::jsonb) || EXCLUDED.metadata,
                    status = 'active',
                    updated_at = NOW()
                RETURNING item_id::text
                """,
                (
                    pipeline,
                    machine_category,
                    machine,
                    problem_category,
                    problem,
                    root_cause,
                    solution,
                    json.dumps(payload.metadata or {}, ensure_ascii=False),
                ),
            )
            item_id = (cur.fetchone() or [None])[0]
            cur.execute(
                """
                INSERT INTO knowledge_item_weights (item_id, helpful_weight, misleading_weight, feedback_count, current_weight)
                VALUES (%s::uuid, 0, 0, 0, 0.5)
                ON CONFLICT (item_id) DO NOTHING
                """,
                (item_id,),
            )
            conn.commit()

    return {"item_id": item_id}


@router.get("/items")
async def list_knowledge_items(pipeline: str | None = None, status: str = "active", limit: int = 100, offset: int = 0, user: dict = Depends(require_expert)):
    limit = max(1, min(int(limit or 100), 500))
    offset = max(0, int(offset or 0))
    pipeline_value = _normalize_pipeline(pipeline) if pipeline else None
    status = (status or "active").strip()

    filters = ["ki.status = %s"]
    params: list = [status]
    if pipeline_value:
        filters.append("ki.pipeline = %s")
        params.append(pipeline_value)
    where_sql = " AND ".join(filters)

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                f"""
                SELECT
                    ki.item_id::text,
                    ki.pipeline,
                    ki.machine_category,
                    ki.machine,
                    ki.problem_category,
                    ki.problem,
                    ki.root_cause,
                    ki.solution,
                    ki.metadata,
                    ki.status,
                    ki.created_at,
                    ki.updated_at,
                    COALESCE(kw.helpful_weight, 0),
                    COALESCE(kw.misleading_weight, 0),
                    COALESCE(kw.feedback_count, 0),
                    COALESCE(kw.current_weight, 0.5),
                    kw.expert_weight,
                    COALESCE(kw.expert_weight, kw.current_weight, 0.5) AS effective_weight
                FROM knowledge_items ki
                LEFT JOIN knowledge_item_weights kw ON kw.item_id = ki.item_id
                WHERE {where_sql}
                ORDER BY ki.updated_at DESC
                LIMIT %s OFFSET %s
                """,
                (*params, limit, offset),
            )
            rows = cur.fetchall() or []

    out = []
    for r in rows:
        metadata = r[8] if isinstance(r[8], dict) else {}
        machine = str(r[3] or "").strip()
        machine_category = _canonicalize_machine_category(r[2], machine)
        problem_category = _canonicalize_problem_category(str(r[4] or "").strip(), str(r[5] or ""), str(r[6] or ""))
        if not machine:
            machine = _infer_device_from_filename(str((metadata or {}).get("filename") or "")) if metadata else ""
        machine = _normalize_machine_name(machine) or machine
        if not machine_category:
            machine_category = _infer_machine_category_from_machine(machine)
        if not problem_category:
            problem_category = _infer_problem_category_from_problem(str(r[5] or ""))
        out.append(
            {
                "item_id": r[0],
                "pipeline": r[1],
                "machine_category": machine_category,
                "machine": machine,
                "problem_category": problem_category,
                "problem": r[5],
                "root_cause": r[6],
                "solution": r[7],
                "metadata": metadata,
                "status": r[9],
                "created_at": r[10].isoformat() if r[10] else None,
                "updated_at": r[11].isoformat() if r[11] else None,
                "helpful_weight": float(r[12] or 0),
                "misleading_weight": float(r[13] or 0),
                "feedback_count": int(r[14] or 0),
                "current_weight": float(r[15] if r[15] is not None else 0.5),
                "expert_weight": float(r[16]) if r[16] is not None else None,
                "effective_weight": float(r[17] if r[17] is not None else (r[15] if r[15] is not None else 0.5)),
            }
        )
    return out


@router.post("/items/reextract")
async def reextract_knowledge_items(payload: KnowledgeItemReextractRequest, user: dict = Depends(require_expert)):
    pipeline = _normalize_pipeline(payload.pipeline)
    mode = (payload.mode or "replace").strip().lower()
    if mode not in {"replace", "append"}:
        mode = "replace"

    doc_ids = payload.doc_ids or []
    if doc_ids:
        doc_ids = [str(uuid.UUID(str(x))) for x in doc_ids]

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)

        if not doc_ids:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT d.doc_id::text
                    FROM documents d
                    WHERE d.status = 'active'
                      AND COALESCE(d.metadata->>'pipeline', '流水线1') = %s
                    ORDER BY d.updated_at DESC, d.created_at DESC
                    LIMIT 50
                    """,
                    (pipeline,),
                )
                doc_ids = [r[0] for r in (cur.fetchall() or []) if r and r[0]]

        if not doc_ids:
            return {"pipeline": pipeline, "mode": mode, "doc_ids": [], "deleted": 0, "result": None}

        deleted = 0
        if mode == "replace":
            with conn.cursor() as cur:
                cur.execute(
                    """
                    DELETE FROM knowledge_items
                    WHERE pipeline = %s
                      AND COALESCE(metadata->>'doc_id', '') = ANY(%s::text[])
                    """,
                    (pipeline, doc_ids),
                )
                deleted = int(cur.rowcount or 0)
                conn.commit()

    if (settings.LLM_PROVIDER or "").lower() == "minimax" and (not settings.MINIMAX_API_KEY or not settings.MINIMAX_GROUP_ID):
        raise HTTPException(status_code=503, detail="LLM 未配置：缺少 MINIMAX_API_KEY 或 MINIMAX_GROUP_ID，无法整理。")

    try:
        result = await extract_knowledge_items_with_ai(pipeline=pipeline, doc_ids=doc_ids)
    except HTTPException:
        raise
    except Exception as e:
        detail = str(e).strip() or e.__class__.__name__
        raise HTTPException(status_code=500, detail=f"整理失败: {detail}")
    return {
        "pipeline": pipeline,
        "mode": mode,
        "doc_ids": doc_ids,
        "deleted": deleted,
        "result": {
            "extracted": int(getattr(result, "extracted", 0)),
            "inserted": int(getattr(result, "inserted", 0)),
            "skipped": int(getattr(result, "skipped", 0)),
            "provider": str(getattr(result, "provider", "") or ""),
            "errors": list(getattr(result, "errors", []) or []),
        },
    }


@router.post("/items/cleanup")
async def cleanup_knowledge_items(payload: KnowledgeItemCleanupRequest, user: dict = Depends(require_expert)):
    pipeline = _normalize_pipeline(payload.pipeline)
    dry_run = bool(payload.dry_run)
    delete_unknown_cause = bool(payload.delete_unknown_cause)
    delete_noise = bool(payload.delete_noise)

    machine_noise = {
        "操作说明", "使用说明", "用户指南", "用户手册", "说明书", "安装指南", "安装说明",
        "参数表", "技术参数", "目录", "保养计划", "安全须知", "注意事项",
    }
    unknown_cause = {"未明确", "未知", "不详", "n/a", "na", "-", "—", "null", "none"}
    noise_in_problem = {"操作说明", "使用说明", "用户指南", "说明书", "安装", "参数", "目录", "注意", "安全", "步骤", "工具", "保养", "授权", "接触器"}
    machine_noise_norm = {_norm_text(x) for x in machine_noise}
    unknown_cause_norm = {_norm_text(x) for x in unknown_cause}
    noise_problem_norm = {_norm_text(x) for x in noise_in_problem}

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    ki.item_id::text,
                    ki.machine_category,
                    ki.machine,
                    ki.problem_category,
                    ki.problem,
                    ki.root_cause,
                    ki.metadata
                FROM knowledge_items ki
                WHERE ki.status = 'active'
                  AND ki.pipeline = %s
                ORDER BY ki.updated_at DESC
                """,
                (pipeline,),
            )
            rows = cur.fetchall() or []

        to_delete: list[str] = []
        to_update: list[tuple] = []
        for item_id, mc, m, pc, p, rc, meta in rows:
            metadata = meta if isinstance(meta, dict) else {}
            machine = str(m or "").strip()
            machine_category = str(mc or "").strip()
            problem_category = str(pc or "").strip()
            problem = str(p or "").strip()
            root_cause = str(rc or "").strip()

            if not machine:
                machine = _infer_device_from_filename(str((metadata or {}).get("filename") or "")) if metadata else ""
            machine_category = _canonicalize_machine_category(machine_category, machine)
            if not machine_category and machine:
                machine_category = _infer_machine_category_from_machine(machine)
            if not problem_category and problem:
                problem_category = _infer_problem_category_from_problem(problem)

            norm_machine = _norm_text(machine)
            norm_root = _norm_text(root_cause)
            norm_problem = _norm_text(problem)

            should_delete = False
            if delete_unknown_cause and (not root_cause or norm_root in unknown_cause_norm):
                should_delete = True
            if delete_noise:
                if norm_machine in machine_noise_norm:
                    should_delete = True
                if any(k in norm_problem for k in noise_problem_norm if k):
                    should_delete = True
                if _is_noise_phrase(problem):
                    should_delete = True
                if machine_category == "" and (norm_machine in machine_noise_norm or any(k in norm_problem for k in noise_problem_norm if k)):
                    should_delete = True

            if should_delete:
                to_delete.append(str(item_id))
            else:
                to_update.append((machine_category, machine, problem_category, str(item_id)))

        deleted = len(to_delete)
        updated = 0

        if not dry_run:
            with conn.cursor() as cur:
                if to_delete:
                    cur.execute(
                        "UPDATE knowledge_items SET status = 'deleted', updated_at = NOW() WHERE item_id = ANY(%s::uuid[])",
                        (to_delete,),
                    )
                if to_update:
                    cur.executemany(
                        """
                        UPDATE knowledge_items
                        SET machine_category = %s,
                            machine = %s,
                            problem_category = %s,
                            updated_at = NOW()
                        WHERE item_id = %s::uuid
                        """,
                        to_update,
                    )
                    updated = int(cur.rowcount or 0)
                conn.commit()
        else:
            updated = len(to_update)

    return {"pipeline": pipeline, "dry_run": dry_run, "deleted": deleted, "updated": updated, "total": len(rows)}


@router.post("/items/autofill")
async def autofill_knowledge_items(payload: KnowledgeItemAutoFillRequest, user: dict = Depends(require_expert)):
    pipeline = _normalize_pipeline(payload.pipeline)
    limit = max(1, min(int(payload.limit or 60), 300))
    dry_run = bool(payload.dry_run)

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    ki.item_id::text,
                    ki.machine_category,
                    ki.machine,
                    ki.problem_category,
                    ki.problem,
                    ki.root_cause,
                    ki.solution,
                    ki.metadata
                FROM knowledge_items ki
                WHERE ki.status = 'active'
                  AND ki.pipeline = %s
                  AND (
                    COALESCE(ki.machine_category, '') = ''
                    OR COALESCE(ki.machine, '') = ''
                    OR COALESCE(ki.problem_category, '') = ''
                    OR COALESCE(ki.problem, '') = ''
                    OR COALESCE(ki.root_cause, '') = ''
                    OR COALESCE(ki.solution, '') = ''
                  )
                ORDER BY ki.updated_at DESC
                LIMIT %s
                """,
                (pipeline, limit),
            )
            rows = cur.fetchall() or []

    updated_rows: list[tuple] = []
    errors: list[dict] = []
    scanned = len(rows)

    for item_id, mc, m, pc, p, rc, sol, meta in rows:
        metadata = meta if isinstance(meta, dict) else {}
        machine = str(m or "").strip()
        machine_category = str(mc or "").strip()
        problem_category = str(pc or "").strip()
        problem = str(p or "").strip()
        root_cause = str(rc or "").strip()
        solution = str(sol or "").strip()

        if not machine:
            machine = _infer_device_from_filename(str((metadata or {}).get("filename") or "")) if metadata else ""

        if not machine_category:
            machine_category = _infer_machine_category_from_machine(machine) if machine else _infer_machine_category_from_machine(root_cause)
        if not problem_category and problem:
            problem_category = _infer_problem_category_from_problem(problem)

        need_ai = (not machine_category) or (not problem_category) or (not machine) or (not problem) or (not root_cause) or (not solution)
        if need_ai:
            try:
                ai_out = await _ai_fill_item_blanks(pipeline, machine_category, machine, problem_category, problem, root_cause, solution, metadata)
                if not machine and ai_out.get("machine"):
                    machine = ai_out.get("machine", "")
                if not problem and ai_out.get("problem"):
                    problem = ai_out.get("problem", "")
                if not root_cause and ai_out.get("root_cause"):
                    root_cause = ai_out.get("root_cause", "")
                if not solution and ai_out.get("solution"):
                    solution = ai_out.get("solution", "")

                if not machine_category and ai_out.get("machine_category"):
                    machine_category = ai_out.get("machine_category", "")
                if not problem_category and ai_out.get("problem_category"):
                    problem_category = ai_out.get("problem_category", "")

                if (not machine_category) or (not problem_category):
                    try:
                        ai_cat = await _ai_enrich_categories(pipeline, machine, problem, root_cause)
                        if not machine_category:
                            machine_category = ai_cat.get("machine_category", "") or machine_category
                        if not problem_category:
                            problem_category = ai_cat.get("problem_category", "") or problem_category
                    except Exception:
                        pass
            except Exception as e:
                errors.append({"item_id": str(item_id), "error": str(e) or e.__class__.__name__})
                continue

        machine_category = _canonicalize_machine_category(machine_category, machine)
        if not machine_category and machine:
            machine_category = _infer_machine_category_from_machine(machine)
        if not problem_category and problem:
            problem_category = _infer_problem_category_from_problem(problem)

        machine = _clean_single_line_text(machine, 160)
        machine_category = _clean_single_line_text(machine_category, 120)
        problem_category = _clean_single_line_text(problem_category, 120)
        problem = _clean_single_line_text(problem, 240)
        root_cause = _clean_single_line_text(root_cause, 320)
        solution = _clean_single_line_text(solution, 800)

        updated_rows.append((machine_category, machine, problem_category, problem, root_cause, solution, str(item_id)))

    updated = 0
    if not dry_run and updated_rows:
        with psycopg2.connect(
            host=settings.DB_HOST, port=settings.DB_PORT,
            user=settings.DB_USER, password=settings.DB_PASSWORD,
            database=settings.DB_NAME
        ) as conn:
            _ensure_structured_knowledge_tables(conn)
            with conn.cursor() as cur:
                cur.executemany(
                    """
                    UPDATE knowledge_items
                    SET machine_category = %s,
                        machine = %s,
                        problem_category = %s,
                        problem = %s,
                        root_cause = %s,
                        solution = %s,
                        updated_at = NOW()
                    WHERE item_id = %s::uuid
                    """,
                    updated_rows,
                )
                updated = int(cur.rowcount or 0)
                conn.commit()
    else:
        updated = len(updated_rows)

    return {
        "pipeline": pipeline,
        "dry_run": dry_run,
        "scanned": scanned,
        "updated": updated,
        "errors": errors[:50],
    }



@router.put("/items/{item_id}")
async def update_knowledge_item(item_id: str, payload: KnowledgeItemUpdateRequest, user: dict = Depends(require_expert)):
    try:
        item_uuid = uuid.UUID(str(item_id))
    except Exception:
        raise HTTPException(status_code=400, detail="item_id 不是有效的 UUID")

    fields = []
    params: list = []
    for col in ["pipeline", "machine_category", "machine", "problem_category", "problem", "root_cause", "solution", "status"]:
        v = getattr(payload, col)
        if v is None:
            continue
        if col == "pipeline":
            v = _normalize_pipeline(v)
        if col == "machine_category":
            v = _canonicalize_machine_category(v, payload.machine)
        fields.append(f"{col} = %s")
        params.append(v)
    if payload.metadata is not None:
        fields.append("metadata = %s::jsonb")
        params.append(json.dumps(payload.metadata or {}, ensure_ascii=False))
    fields.append("updated_at = NOW()")

    if not fields:
        return {"item_id": str(item_uuid)}

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE knowledge_items SET {', '.join(fields)} WHERE item_id = %s",
                (*params, str(item_uuid)),
            )
            conn.commit()

    return {"item_id": str(item_uuid)}


@router.delete("/items/{item_id}")
async def delete_knowledge_item(item_id: str, user: dict = Depends(require_expert)):
    try:
        item_uuid = uuid.UUID(str(item_id))
    except Exception:
        raise HTTPException(status_code=400, detail="item_id 不是有效的 UUID")

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute("UPDATE knowledge_items SET status='deleted', updated_at=NOW() WHERE item_id = %s", (str(item_uuid),))
            conn.commit()

    return {"item_id": str(item_uuid), "deleted": True}


@router.post("/items/search")
async def search_knowledge_items(payload: KnowledgeItemSearchRequest):
    query = str(payload.query or "").strip()
    top_k = max(1, min(int(payload.top_k or 10), 50))
    pipeline_value = _normalize_pipeline(payload.pipeline) if payload.pipeline else None

    if not query:
        return {"results": [], "count": 0, "query": query, "pipeline": pipeline_value or ""}

    filters = ["ki.status = 'active'"]
    params: list = []
    if pipeline_value:
        filters.append("ki.pipeline = %s")
        params.append(pipeline_value)
    where_sql = " AND ".join(filters)

    like = f"%{query}%"
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                f"""
                SELECT
                    ki.item_id::text,
                    ki.pipeline,
                    ki.machine_category,
                    ki.machine,
                    ki.problem_category,
                    ki.problem,
                    ki.root_cause,
                    ki.solution,
                    COALESCE(kw.expert_weight, kw.current_weight, 0.5) AS item_weight
                FROM knowledge_items ki
                LEFT JOIN knowledge_item_weights kw ON kw.item_id = ki.item_id
                WHERE {where_sql}
                  AND (ki.problem ILIKE %s OR ki.root_cause ILIKE %s OR ki.solution ILIKE %s)
                ORDER BY item_weight DESC, ki.updated_at DESC
                LIMIT %s
                """,
                (*params, like, like, like, top_k),
            )
            rows = cur.fetchall() or []

    results = []
    for r in rows:
        results.append(
            {
                "item_id": r[0],
                "pipeline": r[1],
                "machine_category": _canonicalize_machine_category(r[2], r[3]),
                "machine": r[3],
                "problem_category": r[4],
                "problem": r[5],
                "root_cause": r[6],
                "solution": r[7],
                "item_weight": float(r[8] if r[8] is not None else 0.5),
            }
        )
    return {"results": results, "count": len(results), "query": query, "pipeline": pipeline_value or ""}


@router.post("/items/feedback-weight")
async def feedback_knowledge_item_weight(payload: KnowledgeItemWeightFeedbackRequest):
    try:
        item_uuid = uuid.UUID(str(payload.item_id))
    except Exception:
        raise HTTPException(status_code=400, detail="item_id 不是有效的 UUID")

    feedback_type = str(payload.feedback_type or "").strip()
    if feedback_type not in {"helpful", "misleading"}:
        raise HTTPException(status_code=400, detail="feedback_type 仅支持 helpful / misleading")
    amount = float(payload.amount or 1.0)
    if amount <= 0:
        amount = 1.0

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO knowledge_item_weights (item_id, helpful_weight, misleading_weight, feedback_count, current_weight)
                VALUES (%s, 0, 0, 0, 0.5)
                ON CONFLICT (item_id) DO NOTHING
                """,
                (str(item_uuid),),
            )
            cur.execute(
                "SELECT helpful_weight, misleading_weight FROM knowledge_item_weights WHERE item_id = %s",
                (str(item_uuid),),
            )
            row = cur.fetchone() or (0, 0)
            helpful = float(row[0] or 0)
            misleading = float(row[1] or 0)
            if feedback_type == "helpful":
                helpful += amount
            else:
                misleading += amount
            total = max(1.0, helpful + misleading)
            weight = helpful / total
            weight = max(0.0, min(1.0, weight))

            cur.execute(
                """
                UPDATE knowledge_item_weights
                SET helpful_weight = %s,
                    misleading_weight = %s,
                    feedback_count = feedback_count + 1,
                    current_weight = %s,
                    updated_at = NOW()
                WHERE item_id = %s
                """,
                (helpful, misleading, weight, str(item_uuid)),
            )
            cur.execute(
                "SELECT COALESCE(expert_weight, %s) FROM knowledge_item_weights WHERE item_id = %s",
                (weight, str(item_uuid)),
            )
            effective_weight = float((cur.fetchone() or [weight])[0] or weight)
            conn.commit()

    return {"item_id": str(item_uuid), "item_weight": weight, "effective_weight": effective_weight}


@router.post("/items/expert-weight")
async def set_knowledge_item_expert_weight(payload: KnowledgeItemExpertWeightRequest, user: dict = Depends(require_expert)):
    try:
        item_uuid = uuid.UUID(str(payload.item_id))
    except Exception:
        raise HTTPException(status_code=400, detail="item_id 不是有效的 UUID")

    w = payload.expert_weight
    if w is not None:
        w = float(w)
        if w > 1.0:
            w = w / 100.0
        if w < 0.0 or w > 1.0:
            raise HTTPException(status_code=400, detail="expert_weight 取值范围为 0~1（或 0~100）")
        w = round(w, 4)

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO knowledge_item_weights (item_id, helpful_weight, misleading_weight, feedback_count, current_weight, expert_weight)
                VALUES (%s, 0, 0, 0, 0.5, %s)
                ON CONFLICT (item_id) DO UPDATE SET expert_weight = EXCLUDED.expert_weight, updated_at = NOW()
                """,
                (str(item_uuid), w),
            )
            cur.execute(
                "SELECT COALESCE(expert_weight, current_weight, 0.5), current_weight, expert_weight FROM knowledge_item_weights WHERE item_id = %s",
                (str(item_uuid),),
            )
            row = cur.fetchone() or (0.5, 0.5, None)
            conn.commit()

    return {"item_id": str(item_uuid), "effective_weight": float(row[0]), "current_weight": float(row[1]), "expert_weight": (float(row[2]) if row[2] is not None else None)}


@router.get("/items/suggestions")
async def list_knowledge_item_suggestions(pipeline: str = "流水线1", limit: int = 8):
    pipeline = _normalize_pipeline(pipeline)
    limit = max(1, min(int(limit or 8), 20))

    def sanitize_fault_suggestion(raw: str, max_len: int = 24) -> str:
        src = str(raw or "").strip()
        if not src:
            return ""
        src = re.sub(r"[\r\n\t]+", " ", src)
        src = re.sub(r"\s+", " ", src).strip()
        compact = re.sub(r"\s+", "", src)
        if len(compact) < 4:
            return ""
        unsafe = ["人员伤亡", "危险电压", "触电", "警告", "安全须知", "注意事项", "操作说明", "使用说明", "说明书", "安装指南", "安装说明", "错误操作"]
        if any(k in compact for k in unsafe):
            return ""
        has_fault = any(k in compact for k in (FAULT_HINTS + ["无法", "不能", "不亮", "不动作", "不通电", "过温", "误报警", "动作缓慢"]))
        if not has_fault:
            return ""
        if any(k in compact for k in SOLUTION_HINTS) and ("无法" not in compact):
            return ""
        cleaned = re.sub(r"[^0-9a-zA-Z\u4e00-\u9fff ]+", " ", src)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if not cleaned:
            return ""
        normalized = cleaned.replace(" ", "") if re.search(r"[\u4e00-\u9fff]", cleaned) else cleaned
        return normalized[:max_len].strip()

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    ki.problem,
                    COALESCE(kw.expert_weight, kw.current_weight, 0.5) AS w,
                    ki.updated_at
                FROM knowledge_items ki
                LEFT JOIN knowledge_item_weights kw ON kw.item_id = ki.item_id
                WHERE ki.status = 'active'
                  AND ki.pipeline = %s
                  AND LENGTH(TRIM(COALESCE(ki.problem, ''))) > 0
                ORDER BY w DESC, ki.updated_at DESC
                LIMIT 200
                """,
                (pipeline,),
            )
            rows = cur.fetchall() or []

    seen = set()
    out = []
    for problem, _, _ in rows:
        p = sanitize_fault_suggestion(problem)
        if not p:
            continue
        key = p.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(p)
        if len(out) >= limit:
            break
    return {"pipeline": pipeline, "suggestions": out}


def _extract_manual_entries(pipeline: str, limit: int = 400) -> list[dict]:
    pipeline = _normalize_pipeline(pipeline)
    limit = max(50, min(int(limit or 400), 1200))

    safety_keys = [
        "安全", "警告", "注意", "危险", "触电", "高压", "电压", "防护", "佩戴", "护目镜", "手套",
        "断电", "锁定挂牌", "防火", "爆炸", "烫伤", "烧伤", "夹伤", "挤压", "坠落", "泄漏",
    ]
    op_keys = [
        "操作", "步骤", "流程", "启动", "停止", "开机", "关机", "复位", "设置", "调试", "校准",
        "安装", "拆卸", "更换", "清理", "连接", "按下", "打开", "关闭", "检查", "确认",
    ]
    norm_keys = [
        "规范", "要求", "标准", "应", "必须", "不得", "严禁", "禁止", "允许", "范围", "参数", "扭矩",
        "力矩", "公差", "验收", "检验", "维护周期", "保养周期",
    ]
    unsafe = [
        "目录", "更新日期", "查询热线", "本手册", "官方完整手册", "标准化模板", "产品结构图示",
        "技术参数", "参数表",
    ]

    def norm_line(s: str) -> str:
        v = str(s or "").strip()
        v = re.sub(r"[\r\n\t]+", " ", v)
        v = re.sub(r"\s+", " ", v).strip()
        v = re.sub(r"^[\s·\-•\d\.\、]+", "", v).strip()
        v = re.sub(r"[^0-9a-zA-Z\u4e00-\u9fff ]+", " ", v)
        v = re.sub(r"\s+", " ", v).strip()
        return v

    def pick_category(compact: str) -> str:
        if any(k in compact for k in safety_keys):
            return "安全"
        if any(k in compact for k in op_keys):
            return "操作"
        if any(k in compact for k in norm_keys):
            return "规范"
        return ""

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    d.doc_id::text,
                    d.filename,
                    COALESCE(c.page_num, 0) AS page_num,
                    COALESCE(c.chunk_index, 0) AS chunk_index,
                    c.text
                FROM documents d
                JOIN document_chunks c ON c.doc_id = d.doc_id
                WHERE d.status = 'active'
                  AND COALESCE(d.metadata->>'pipeline', '流水线1') = %s
                  AND LENGTH(TRIM(COALESCE(c.text, ''))) > 0
                ORDER BY d.upload_time DESC, c.chunk_index ASC
                LIMIT 2000
                """,
                (pipeline,),
            )
            rows = cur.fetchall() or []

    out: list[dict] = []
    seen = set()
    for doc_id, filename, page_num, chunk_index, text in rows:
        raw = str(text or "")
        for part in re.split(r"[\n\r]+", raw):
            v = norm_line(part)
            if not v:
                continue
            compact = re.sub(r"\s+", "", v)
            if len(compact) < 6 or len(compact) > 120:
                continue
            if any(k in compact for k in unsafe):
                continue
            if len(compact) <= 10 and (_is_noise_phrase(v) or _is_noise_phrase(compact)):
                continue
            cat = pick_category(compact)
            if not cat:
                continue
            k = (cat, _norm_text(compact))
            if k in seen:
                continue
            seen.add(k)
            out.append(
                {
                    "category": cat,
                    "text": v[:240],
                    "source": {
                        "doc_id": doc_id,
                        "filename": filename,
                        "page": int(page_num or 0),
                        "chunk_index": int(chunk_index or 0),
                    },
                }
            )
            if len(out) >= limit:
                return out
    return out


@router.get("/manual/entries")
async def list_manual_entries(pipeline: str = "流水线1", category: str | None = None, limit: int = 400, use_ai: bool = True):
    pipeline = _normalize_pipeline(pipeline)
    cat = str(category or "").strip()
    if use_ai:
        manual = await _build_manual_with_ai(pipeline)
        entries = _flatten_manual_entries(manual)
    else:
        entries = [{"category": x.get("category"), "topic": "", "text": x.get("text")} for x in _extract_manual_entries(pipeline, limit=limit)]
    if cat:
        entries = [x for x in entries if str((x or {}).get("category") or "") == cat]
    return {"pipeline": pipeline, "category": cat, "count": len(entries), "entries": entries}


@router.get("/manual/export/word")
async def export_manual_word(pipeline: str = "流水线1", use_ai: bool = True):
    pipeline = _normalize_pipeline(pipeline)
    if use_ai:
        manual = await _build_manual_with_ai(pipeline)
    else:
        manual = {"pipeline": pipeline, "sections": _group_fallback_manual(_extract_manual_entries(pipeline, limit=1200))}

    doc = Document()
    doc.add_heading("规范手册", 0)
    doc.add_paragraph(f"流水线：{pipeline}")
    doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    for sec in (manual or {}).get("sections", []) or []:
        cat = str((sec or {}).get("category") or "").strip()
        title = "安全规范" if cat == "安全" else "操作规范" if cat == "操作" else "通用规范"
        doc.add_heading(title, level=1)
        topics = (sec or {}).get("topics", []) or []
        if not topics:
            doc.add_paragraph("暂无条目")
            doc.add_paragraph("")
            continue
        for t in topics:
            t_title = str((t or {}).get("title") or "").strip()
            if t_title:
                doc.add_heading(t_title, level=2)
            items = (t or {}).get("items", []) or []
            for it in items:
                v = str(it or "").strip()
                if v:
                    doc.add_paragraph(v[:320], style="List Bullet")
        doc.add_paragraph("")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    filename = f"规范手册_{pipeline}.docx"
    return FileResponse(
        tmp.name,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _group_fallback_manual(entries: list[dict]) -> list[dict]:
    grouped: dict[str, list[str]] = {"安全": [], "操作": [], "规范": []}
    for e in entries or []:
        c = str((e or {}).get("category") or "").strip()
        t = str((e or {}).get("text") or "").strip()
        if c in grouped and t:
            grouped[c].append(t)
    sections = []
    for cat in ["安全", "操作", "规范"]:
        items = grouped.get(cat) or []
        sections.append({"category": cat, "topics": [{"title": "", "items": items[:400]}]})
    return sections


def _flatten_manual_entries(manual: dict) -> list[dict]:
    out: list[dict] = []
    for sec in (manual or {}).get("sections", []) or []:
        cat = str((sec or {}).get("category") or "").strip()
        for t in (sec or {}).get("topics", []) or []:
            topic = str((t or {}).get("title") or "").strip()
            for it in (t or {}).get("items", []) or []:
                text = str(it or "").strip()
                if cat and text:
                    out.append({"category": cat, "topic": topic, "text": text})
    return out


_MANUAL_CACHE = {}
_MANUAL_CACHE_TTL = 3600  # 1 hour

async def _build_manual_with_ai(pipeline: str) -> dict:
    global _MANUAL_CACHE
    now = time.time()
    if pipeline in _MANUAL_CACHE:
        cached_data, timestamp = _MANUAL_CACHE[pipeline]
        if now - timestamp < _MANUAL_CACHE_TTL:
            return cached_data

    pipeline = _normalize_pipeline(pipeline)
    base = _extract_manual_entries(pipeline, limit=420)
    grouped: dict[str, list[str]] = {"安全": [], "操作": [], "规范": []}
    for e in base:
        cat = str((e or {}).get("category") or "").strip()
        text = str((e or {}).get("text") or "").strip()
        if cat in grouped and text:
            grouped[cat].append(text)
    safety = "\n".join(f"- {x}" for x in grouped["安全"][:80])
    operation = "\n".join(f"- {x}" for x in grouped["操作"][:80])
    standard = "\n".join(f"- {x}" for x in grouped["规范"][:80])
    safety = safety[:2500]
    operation = operation[:2500]
    standard = standard[:2500]

    prompt = f"""
你是工业现场“规范手册”编写助手。请基于给定的材料，整理输出一份更详细、更规范、可直接执行的《规范手册》。

要求
1) 严格只输出 JSON，不要任何解释文字
2) 手册按三个大章输出：安全、操作、规范
3) 每章必须包含 3~6 个主题（topics），每个主题包含 3~5 条可执行条目（items）
4) items 必须是“动作/约束句”，如：必须/不得/应/严禁/确认/佩戴/断电/锁定挂牌/检查/校准/记录
5) 去掉来源信息：不得出现“来源/文件名/页码/doc_id/第X页”等字样
6) 去掉空泛重复：合并同义条目，避免重复句式
7) 允许合理补全细节，但不得编造具体型号/数值/标准编号；若需要数值，写成“按现场标准/按设备铭牌/按工艺要求”

输出结构（严格遵循）：
{{
  "pipeline": "{pipeline}",
  "sections": [
    {{
      "category": "安全",
      "topics": [
        {{
          "title": "主题标题",
          "items": ["条目1", "条目2"]
        }}
      ]
    }},
    {{
      "category": "操作",
      "topics": [{{"title":"主题标题","items":["条目1"]}}]
    }},
    {{
      "category": "规范",
      "topics": [{{"title":"主题标题","items":["条目1"]}}]
    }}
  ]
}}

材料（安全）：
{safety}

材料（操作）：
{operation}

材料（规范）：
{standard}
"""
    try:
        manager = get_llm_manager()
        resp, _provider = await manager.generate_with_fallback(prompt)
        parsed = _extract_json_dict(resp.content if resp else "")
        sections = parsed.get("sections", []) if isinstance(parsed, dict) else []
        cleaned_sections = []
        allowed = {"安全", "操作", "规范"}
        for sec in (sections or [])[:3]:
            cat = str((sec or {}).get("category") or "").strip()
            if cat not in allowed:
                continue
            topics = []
            for t in (sec or {}).get("topics", []) or []:
                title = str((t or {}).get("title") or "").strip()
                items = []
                for it in (t or {}).get("items", []) or []:
                    v = str(it or "").strip()
                    if not v:
                        continue
                    if any(k in v for k in ["来源", "文件名", "页", "doc_id", "chunk"]):
                        continue
                    items.append(v[:120])
                if items:
                    topics.append({"title": title[:30], "items": items[:10]})
                if len(topics) >= 12:
                    break
            cleaned_sections.append({"category": cat, "topics": topics[:12]})

        if cleaned_sections:
            res = {"pipeline": pipeline, "sections": cleaned_sections}
            _MANUAL_CACHE[pipeline] = (res, time.time())
            return res
    except Exception:
        pass

    res = {"pipeline": pipeline, "sections": _group_fallback_manual(base)}
    _MANUAL_CACHE[pipeline] = (res, time.time())
    return res


@router.get("/graph")
async def knowledge_graph(pipeline: str = "流水线1"):
    pipeline = _normalize_pipeline(pipeline)

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        _ensure_structured_knowledge_tables(conn)
        try:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT
                        TRIM(COALESCE(machine_category, '')) AS machine_category,
                        COALESCE(NULLIF(TRIM(machine), ''), '设备') AS machine,
                        TRIM(COALESCE(problem_category, '')) AS problem_category,
                        TRIM(COALESCE(problem, '')) AS problem,
                        TRIM(COALESCE(root_cause, '')) AS root_cause,
                        TRIM(COALESCE(solution, '')) AS solution
                    FROM knowledge_items
                    WHERE status = 'active'
                      AND pipeline = %s
                      AND LENGTH(TRIM(COALESCE(problem, ''))) > 0
                    ORDER BY updated_at DESC
                    LIMIT 800
                    """,
                    (pipeline,),
                )
                item_rows = cur.fetchall() or []
        except Exception:
            item_rows = []

        if item_rows:
            unknown_root = {"未明确", "未知", "不详", "n/a", "na", "-", "—", "null", "none"}
            unknown_root_norm = {_norm_text(x) for x in unknown_root}
            machine_noise_norm = {_norm_text(x) for x in {"操作说明", "使用说明", "用户指南", "用户手册", "说明书", "安装指南", "安装说明", "参数表", "技术参数", "目录", "保养计划", "安全须知", "注意事项"}}

            cat_map: dict[str, dict] = {}
            device_map: dict[str, dict] = {}

            for machine_category, machine, problem_category, problem, root_cause, solution in item_rows:
                dev = _normalize_machine_name(str(machine or "").strip()) or (str(machine or "").strip() or "设备")
                dev_norm = _norm_text(dev)
                if not dev_norm or dev_norm in machine_noise_norm:
                    continue
                fault = str(problem or "").strip()
                if not fault or _is_noise_phrase(fault):
                    continue
                root = str(root_cause or "").strip()
                sol = str(solution or "").strip()
                if not root or _norm_text(root) in unknown_root_norm:
                    root = sol
                if not root or _norm_text(root) in unknown_root_norm:
                    continue

                mc = str(machine_category or "").strip()
                if not mc:
                    mc = _infer_machine_category_from_machine(dev)
                mc = _canonicalize_machine_category(mc, dev)
                pc = str(problem_category or "").strip()
                if not pc:
                    pc = _infer_problem_category_from_problem(fault)

                mc_key = _norm_text(mc) or mc.lower()
                m_key = _machine_dedupe_key(dev)
                pc_key = _norm_text(pc) or pc.lower()
                p_key = _norm_text(fault) or fault.lower()
                r_key = _norm_text(root) or root.lower()

                mc_bucket = cat_map.setdefault(mc_key, {"name": mc or "通用设备", "machines": {}})
                mach_bucket = mc_bucket["machines"].setdefault(m_key, {"names": set(), "name": dev, "problem_categories": {}})
                mach_bucket["names"].add(dev)
                mach_bucket["name"] = _pick_preferred_machine_name(mach_bucket["names"])
                pc_bucket = mach_bucket["problem_categories"].setdefault(pc_key, {"name": pc or "", "problems": {}})
                if len(pc) > len(str(pc_bucket.get("name", ""))):
                    pc_bucket["name"] = pc
                prob_bucket = pc_bucket["problems"].setdefault(p_key, {"name": fault, "root_causes": {}})
                if len(fault) > len(str(prob_bucket.get("name", ""))):
                    prob_bucket["name"] = fault
                prob_bucket["root_causes"][r_key] = root

                dev_bucket = device_map.setdefault(m_key, {"names": set(), "name": dev, "faults": {}})
                dev_bucket["names"].add(dev)
                dev_bucket["name"] = _pick_preferred_machine_name(dev_bucket["names"])
                dev_fault = dev_bucket["faults"].setdefault(p_key, {"name": fault, "solutions": []})
                if root not in dev_fault["solutions"]:
                    dev_fault["solutions"].append(root)

            machine_categories = []
            for _k, mc_bucket in sorted(cat_map.items(), key=lambda x: x[0]):
                machines = []
                for _mk, mach_bucket in list((mc_bucket.get("machines") or {}).items())[:30]:
                    pcs = []
                    for _pck, pc_bucket in list((mach_bucket.get("problem_categories") or {}).items())[:30]:
                        problems = []
                        for _pk, p_bucket in list((pc_bucket.get("problems") or {}).items())[:30]:
                            roots = list((p_bucket.get("root_causes") or {}).values())
                            roots = [r for r in roots if str(r).strip()]
                            if not roots:
                                continue
                            problems.append({"name": p_bucket.get("name") or "", "root_causes": roots[:8]})
                        if not problems:
                            continue
                        pcs.append({"name": pc_bucket.get("name") or "", "problems": problems[:12]})
                    if not pcs:
                        continue
                    machines.append({"name": mach_bucket.get("name") or "", "problem_categories": pcs[:12]})
                if not machines:
                    continue
                machine_categories.append({"name": mc_bucket.get("name") or "", "machines": machines[:12]})

            devices = []
            for _dev_key, val in list(device_map.items())[:40]:
                faults = []
                dev_name = str(val.get("name") or "").strip() or "设备"
                for _fk, fentry in list((val.get("faults") or {}).items())[:60]:
                    sols = [s for s in (fentry.get("solutions") or []) if str(s).strip()]
                    if not sols:
                        continue
                    faults.append({"name": str(fentry.get("name") or "").strip(), "solutions": sols[:6]})
                if faults:
                    devices.append({"name": dev_name, "faults": faults[:10]})

            devices = devices[:20]
            return {
                "pipeline": pipeline,
                "doc_count": 0,
                "item_count": len(item_rows),
                "devices": devices,
                "kb_tree": {"machine_categories": machine_categories[:12]},
                "device_count": len(devices),
                "fault_count": sum(len((d or {}).get("faults", []) or []) for d in devices),
                "version": "kg_v6_items_tree",
                "source": "knowledge_items",
            }

        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    d.doc_id,
                    d.filename,
                    COALESCE(d.metadata->>'pipeline', '流水线1') AS pipeline,
                    d.metadata,
                    COALESCE(string_agg(c.text, ' ' ORDER BY c.chunk_index), '') AS content
                FROM documents d
                LEFT JOIN document_chunks c ON c.doc_id = d.doc_id
                WHERE d.status = 'active'
                  AND COALESCE(d.metadata->>'pipeline', '流水线1') = %s
                GROUP BY d.doc_id, d.filename, d.metadata
                ORDER BY d.upload_time DESC
                """,
                (pipeline,),
            )
            rows = cur.fetchall()

    merged = {}
    for row in rows:
        metadata = row[3] or {}
        content = row[4] or ""
        part_from_db = (metadata.get("kg", {}) if isinstance(metadata, dict) else {}) or {}
        part_from_db = _sanitize_kg_payload(part_from_db)
        # 图谱查询时优先基于“当前文档内容”重构，避免旧 metadata.kg 噪声长期残留
        part_from_fallback = _build_fallback_kg(content[:120000])
        part_from_min = _build_minimum_kg(content[:120000])
        if _has_useful_kg(part_from_fallback):
            source = part_from_fallback
        elif _has_useful_kg(part_from_db):
            source = part_from_db
        else:
            source = part_from_min
        source = _constrain_kg_to_main_device(source, content[:120000], row[1] or "")
        if not source.get("devices"):
            guessed = _infer_main_device((row[1] or "") + "\n" + content[:1000])
            source = {
                "devices": [
                    {
                        "name": guessed,
                        "faults": [
                            {
                                "name": "设备运行异常",
                                "solutions": ["请补充包含“故障现象-解决方法”的手册段落后重建图谱"]
                            }
                        ]
                    }
                ]
            }
        part = {}
        for d in source.get("devices", []):
            dev_name = str((d or {}).get("name", "")).strip()
            if not dev_name:
                continue
            part.setdefault(dev_name, {"faults": {}})
            for f in (d or {}).get("faults", []) or []:
                fn = str((f or {}).get("name", "")).strip()
                if not fn:
                    continue
                part[dev_name]["faults"].setdefault(fn, {"solutions": []})
                for s in (f or {}).get("solutions", []) or []:
                    sv = str(s).strip()
                    if sv and sv not in part[dev_name]["faults"][fn]["solutions"]:
                        part[dev_name]["faults"][fn]["solutions"].append(sv)
        for dev, val in part.items():
            merged.setdefault(dev, {"faults": {}})
            for fault, fval in val["faults"].items():
                merged[dev]["faults"].setdefault(fault, {"solutions": []})
                for sol in fval["solutions"]:
                    if sol not in merged[dev]["faults"][fault]["solutions"]:
                        merged[dev]["faults"][fault]["solutions"].append(sol)

    devices = []
    for dev, val in merged.items():
        faults = []
        for fault, fval in val["faults"].items():
            if not _is_valid_fault_phrase(fault):
                continue
            clean_solutions = []
            seen_sol = set()
            for s in fval["solutions"]:
                sv = _clean_solution(str(s))
                nk = _norm_text(sv)
                if sv and nk not in seen_sol:
                    clean_solutions.append(sv)
                    seen_sol.add(nk)
            if not clean_solutions:
                continue
            faults.append({
                "name": fault,
                "solutions": clean_solutions[:5],
            })
        if not faults:
            continue
        devices.append({
            "name": dev,
            "faults": faults[:8],
        })

    if not devices and rows:
        guessed = _infer_main_device((rows[0][1] or "") + "\n" + (rows[0][4] or "")[:1000])
        devices = [{
            "name": guessed,
            "faults": [{
                "name": "设备运行异常",
                "solutions": ["检查电源与关键部件连接状态后重建图谱"]
            }]
        }]

    return {
        "pipeline": pipeline,
        "doc_count": len(rows),
        "devices": devices[:20],
        "device_count": len(devices[:20]),
        "fault_count": sum(len((d or {}).get("faults", []) or []) for d in devices[:20]),
        "version": "kg_v4_content_first",
    }


@router.post("/graph/rebuild")
async def rebuild_knowledge_graph(pipeline: str = "流水线1"):
    pipeline = _normalize_pipeline(pipeline)
    rebuilt = 0
    failed = 0

    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT
                    d.doc_id,
                    COALESCE(string_agg(c.text, E'\n' ORDER BY c.chunk_index), '') AS content
                FROM documents d
                LEFT JOIN document_chunks c ON c.doc_id = d.doc_id
                WHERE d.status = 'active'
                  AND COALESCE(d.metadata->>'pipeline', '流水线1') = %s
                GROUP BY d.doc_id
                """,
                (pipeline,),
            )
            rows = cur.fetchall()

    for row in rows:
        doc_id = str(row[0])
        content = row[1] or ""
        try:
            try:
                kg_payload = await _build_doc_kg_with_ai(content[:120000])
                kg_payload = _merge_pairs_into_kg(kg_payload, content[:120000])
                if not _has_useful_kg(kg_payload):
                    kg_payload = _build_fallback_kg(content[:120000])
                if not _has_useful_kg(kg_payload):
                    kg_payload = _build_minimum_kg(content[:120000])
                kg_payload = _constrain_kg_to_main_device(kg_payload, content[:120000])
            except Exception:
                kg_payload = _build_fallback_kg(content[:120000])
                if not _has_useful_kg(kg_payload):
                    kg_payload = _build_minimum_kg(content[:120000])
                kg_payload = _constrain_kg_to_main_device(kg_payload, content[:120000])

            useful = _has_useful_kg(kg_payload)
            with psycopg2.connect(
                host=settings.DB_HOST, port=settings.DB_PORT,
                user=settings.DB_USER, password=settings.DB_PASSWORD,
                database=settings.DB_NAME
            ) as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        """
                        UPDATE documents
                        SET metadata = COALESCE(metadata, '{}'::jsonb) || %s::jsonb
                        WHERE doc_id = %s
                        """,
                        (psycopg2.extras.Json({"kg": kg_payload}), doc_id),
                    )
                    conn.commit()
            if useful:
                rebuilt += 1
            else:
                failed += 1
        except Exception:
            failed += 1

    return {"pipeline": pipeline, "rebuilt": rebuilt, "failed": failed}

@router.post("/reparse")
async def reparse_document(doc_id: str):
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT filename, metadata FROM documents WHERE doc_id=%s AND status <> 'deleted'",
                (doc_id,),
            )
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="文档不存在或已删除")
            meta = row[1] or {}
            path = (meta.get("original_path") if isinstance(meta, dict) else None) or ""
    if not path or not Path(path).exists():
        raise HTTPException(status_code=400, detail="无法定位原文件路径")
    chunks = parse_document(path)
    enc = tiktoken.get_encoding("cl100k_base")
    try:
        for c in chunks:
            c["tokens"] = len(enc.encode(c.get("text", "")))
    except Exception:
        for c in chunks:
            c["tokens"] = None
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM document_chunks WHERE doc_id=%s", (doc_id,))
            conn.commit()
    await add_chunks_to_db(chunks, doc_id)
    merged_text = "\n".join(c.get("text", "") for c in chunks)
    try:
        kg_payload = await _build_doc_kg_with_ai(merged_text)
        kg_payload = _merge_pairs_into_kg(kg_payload, merged_text)
        if not _has_useful_kg(kg_payload):
            kg_payload = _build_fallback_kg(merged_text)
        if not _has_useful_kg(kg_payload):
            kg_payload = _build_minimum_kg(merged_text)
        kg_payload = _constrain_kg_to_main_device(kg_payload, merged_text, row[0] or "")
    except Exception:
        kg_payload = _build_fallback_kg(merged_text)
        if not _has_useful_kg(kg_payload):
            kg_payload = _build_minimum_kg(merged_text)
        kg_payload = _constrain_kg_to_main_device(kg_payload, merged_text, row[0] or "")
    with psycopg2.connect(
        host=settings.DB_HOST, port=settings.DB_PORT,
        user=settings.DB_USER, password=settings.DB_PASSWORD,
        database=settings.DB_NAME
    ) as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE documents SET metadata = COALESCE(metadata,'{}'::jsonb) || %s::jsonb WHERE doc_id=%s",
                (psycopg2.extras.Json({"kg": kg_payload}), doc_id),
            )
            conn.commit()
    return {"doc_id": doc_id, "chunks": len(chunks), "kg_faults": sum(len(f.get("solutions", [])) >= 1 for d in kg_payload.get("devices", []) for f in d.get("faults", []))}
