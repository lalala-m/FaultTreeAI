from fastapi import APIRouter
from fastapi.responses import FileResponse, StreamingResponse
from backend.models.schemas import FaultTree
from docx import Document
from docx.shared import Pt
import tempfile, os
import io

router = APIRouter()


@router.post("/word")
def export_word(fault_tree: FaultTree):
    """导出 Word 报告"""
    doc = Document()
    doc.add_heading("故障树分析报告", 0)

    doc.add_heading("顶事件", level=1)
    doc.add_paragraph(fault_tree.top_event)

    doc.add_heading("分析摘要", level=1)
    doc.add_paragraph(fault_tree.analysis_summary)

    doc.add_heading("节点列表", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "ID", "类型", "名称", "描述"
    for node in fault_tree.nodes:
        row = table.add_row().cells
        row[0].text = node.id
        row[1].text = node.type
        row[2].text = node.name
        row[3].text = node.description

    doc.add_heading("逻辑门", level=1)
    for gate in fault_tree.gates:
        doc.add_paragraph(
            f"{gate.id} [{gate.type}]: {gate.output_node} ← {', '.join(gate.input_nodes)}"
        )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return FileResponse(tmp.name, filename="故障树分析报告.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.post("/pdf")
def export_pdf(fault_tree: FaultTree, mcs: list = None):
    """导出 PDF 报告"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 使用系统默认字体
    from reportlab.pdfbase import pdfmetrics
    from reportlab.lib.pdffonts import Helvetica, HelveticaBold

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    story = []

    # 标题样式
    title_style = ParagraphStyle(
        'Title',
        fontSize=18,
        leading=22,
        alignment=1,
        spaceAfter=20,
    )

    # 副标题样式
    h2_style = ParagraphStyle(
        'H2',
        fontSize=14,
        leading=18,
        spaceBefore=15,
        spaceAfter=10,
    )

    # 正文样式
    body_style = ParagraphStyle(
        'Body',
        fontSize=10,
        leading=14,
        spaceAfter=8,
    )

    # 标题
    story.append(Paragraph("故障树分析报告", title_style))
    story.append(Spacer(1, 0.5 * cm))

    # 顶事件
    story.append(Paragraph("顶事件", h2_style))
    story.append(Paragraph(fault_tree.top_event, body_style))
    story.append(Spacer(1, 0.3 * cm))

    # 分析摘要
    if fault_tree.analysis_summary:
        story.append(Paragraph("分析摘要", h2_style))
        story.append(Paragraph(fault_tree.analysis_summary, body_style))
        story.append(Spacer(1, 0.3 * cm))

    # 节点列表
    story.append(Paragraph("节点列表", h2_style))
    if fault_tree.nodes:
        node_data = [['ID', '类型', '名称', '描述']]
        for node in fault_tree.nodes[:20]:  # 限制最多20行
            node_data.append([
                node.id[:20],
                node.type,
                node.name[:30],
                node.description[:30] if node.description else ''
            ])
        
        node_table = Table(node_data, colWidths=[2*cm, 2*cm, 4*cm, 4*cm])
        node_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
        ]))
        story.append(node_table)
    story.append(Spacer(1, 0.3 * cm))

    # 逻辑门
    story.append(Paragraph("逻辑门", h2_style))
    if fault_tree.gates:
        gate_data = [['门ID', '类型', '输出节点', '输入节点']]
        for gate in fault_tree.gates[:20]:
            gate_data.append([
                gate.id[:15],
                gate.type,
                gate.output_node[:20],
                ', '.join(gate.input_nodes)[:30]
            ])
        
        gate_table = Table(gate_data, colWidths=[2.5*cm, 2*cm, 3*cm, 4.5*cm])
        gate_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
        ]))
        story.append(gate_table)
    story.append(Spacer(1, 0.3 * cm))

    # 最小割集
    if mcs:
        story.append(Paragraph("最小割集（MOCUS）", h2_style))
        for i, cut in enumerate(mcs[:30], 1):  # 限制最多30行
            story.append(Paragraph(f"MCS {i}: {' + '.join(cut)}", body_style))
        if len(mcs) > 30:
            story.append(Paragraph(f"... (共 {len(mcs)} 个最小割集)", body_style))

    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=故障树分析报告.pdf"},
    )
