import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
import re

_CHUNK_SIZE = 500
_CHUNK_OVERLAP = 50
_SEPARATORS = ["\n\n", "\n", "。", "；", " "]


def _merge_with_overlap(chunks: list[str], chunk_size: int, chunk_overlap: int) -> list[str]:
    out: list[str] = []
    buf = ""
    for part in chunks:
        p = (part or "").strip()
        if not p:
            continue
        if not buf:
            buf = p
            continue
        if len(buf) + 1 + len(p) <= chunk_size:
            buf = f"{buf} {p}"
            continue
        out.append(buf.strip())
        overlap = buf[-chunk_overlap:] if chunk_overlap > 0 else ""
        buf = (overlap + " " + p).strip() if overlap else p
    if buf.strip():
        out.append(buf.strip())
    return out


def _recursive_split(text: str, separators: list[str], chunk_size: int) -> list[str]:
    t = str(text or "").strip()
    if not t:
        return []
    if len(t) <= chunk_size:
        return [t]
    if not separators:
        return [t[i : i + chunk_size] for i in range(0, len(t), chunk_size)]
    sep = separators[0]
    parts = t.split(sep) if sep else list(t)
    small: list[str] = []
    for p in parts:
        p2 = p.strip()
        if not p2:
            continue
        if len(p2) <= chunk_size:
            small.append(p2)
        else:
            small.extend(_recursive_split(p2, separators[1:], chunk_size))
    return small


def split_text(text: str) -> list[str]:
    parts = _recursive_split(text, _SEPARATORS, _CHUNK_SIZE)
    return _merge_with_overlap(parts, _CHUNK_SIZE, _CHUNK_OVERLAP)

def _normalize_table_text(text: str) -> str:
    """
    将“常见故障排查表”类表格整理为可解析的文本形式：
    行内多个空白/制表符视为分隔符，生成 '故障现象 | 可能原因 | 解决方法' 结构。
    """
    if not text:
        return ""
    lines = [re.sub(r"[ \t]{2,}", " | ", ln.rstrip()) for ln in text.splitlines()]
    has_header = any(("故障现象" in ln and "可能原因" in ln and "解决方法" in ln) for ln in lines[:6])
    if has_header:
        return "\n".join(lines)
    return text

def parse_pdf(file_path: str) -> list[dict]:
    doc = fitz.open(file_path)
    chunks = []
    engine = None
    for page_num, page in enumerate(doc):
        text = page.get_text("text").strip()
        
        # 如果提取出的文本过少，则认为可能是扫描版图片，尝试进行 OCR 识别
        if len(text) < 50:
            if engine is None:
                try:
                    from rapidocr_onnxruntime import RapidOCR
                    engine = RapidOCR()
                except ImportError:
                    pass
            
            if engine:
                import numpy as np
                import cv2
                # 放大两倍以提高 OCR 精度
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
                if pix.n == 4:
                    img_array = cv2.cvtColor(img_array, cv2.COLOR_RGBA2RGB)
                elif pix.n == 1:
                    img_array = cv2.cvtColor(img_array, cv2.COLOR_GRAY2RGB)
                elif pix.n == 3 and pix.colorspace.name == 'DeviceBGR':
                    img_array = cv2.cvtColor(img_array, cv2.COLOR_BGR2RGB)
                
                result, _ = engine(img_array)
                if result:
                    text = "\n".join([res[1] for res in result])

        # 仅保留 PDF 自带文字层，忽略页面中的图片内容。
        text = _normalize_table_text(text)
        if not text:
            continue
        for i, chunk in enumerate(split_text(text)):
            chunks.append({
                "text": chunk,
                "source": Path(file_path).name,
                "page": page_num + 1,
                "chunk_index": i
            })
    return chunks

def parse_txt(file_path: str) -> list[dict]:
    text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
    text = _normalize_table_text(text)
    chunks = []
    for i, chunk in enumerate(split_text(text)):
        chunks.append({
            "text": chunk,
            "source": Path(file_path).name,
            "page": 0,
            "chunk_index": i
        })
    return chunks

def parse_docx(file_path: str) -> list[dict]:
    doc = Document(file_path)
    paras = []
    for p in doc.paragraphs:
        t = str(p.text or "").strip()
        if not t:
            continue
        if ("→" in t or "->" in t or "⇒" in t or "—>" in t) and not re.match(r"^[-*•]\s+", t):
            t = "- " + t
        paras.append(t)
    full_text = "\n".join(paras)
    # 提取表格并标准化为 'a | b | c' 行
    table_lines = []
    try:
        for tb in doc.tables:
            for row in tb.rows:
                cells = [c.text.strip() for c in row.cells]
                # 合并重复单元格内容
                dedup = []
                for c in cells:
                    if not dedup or c != dedup[-1]:
                        dedup.append(c)
                table_lines.append(" | ".join(dedup))
    except Exception:
        pass
    if table_lines:
        full_text = (full_text + "\n" + "\n".join(table_lines)).strip()
    full_text = _normalize_table_text(full_text)
    chunks = []
    for i, chunk in enumerate(split_text(full_text)):
        chunks.append({
            "text": chunk,
            "source": Path(file_path).name,
            "page": 0,
            "chunk_index": i
        })
    return chunks

def parse_markdown(file_path: str) -> list[dict]:
    """解析 Markdown 文件"""
    text = Path(file_path).read_text(encoding="utf-8")
    chunks = []
    for i, chunk in enumerate(split_text(text)):
        chunks.append({
            "text": chunk,
            "source": Path(file_path).name,
            "page": 0,
            "chunk_index": i
        })
    return chunks

def parse_document(file_path: str) -> list[dict]:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".txt" or ext == ".log":
        return parse_txt(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".md" or ext == ".markdown":
        return parse_markdown(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
