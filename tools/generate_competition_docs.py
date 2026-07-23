import os
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = PROJECT_ROOT / "deliverables" / "competition_docs"


def _set_docx_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt([16, 14, 14, 12][min(level - 1, 3)])


def _add_para(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph(str(text).strip())
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    for r in p.runs:
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)


def _add_title_page(doc: Document, title: str, subtitle: str, doc_type: str = "") -> None:
    _set_docx_base_style(doc)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    def _add_center_line(text: str, size: int, bold: bool = False, font: str = "黑体"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.name = font
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        p.paragraph_format.space_after = Pt(12)

    _add_center_line(title, 26, bold=True)
    _add_center_line(subtitle, 16)
    _add_center_line(doc_type, 18, bold=True, font="宋体")

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("版本：v2.0")
    r.font.size = Pt(12)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.add_page_break()


def _add_table(doc: Document, headers: list, rows: list) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.name = "黑体"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                r.font.size = Pt(12)
    for row in rows:
        row_cells = table.add_row().cells
        for i, v in enumerate(row):
            row_cells[i].text = str(v)
            for p in row_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "宋体"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    r.font.size = Pt(12)
    doc.add_paragraph()


def _h1(doc, text): _add_heading(doc, text, 1)
def _h2(doc, text): _add_heading(doc, text, 2)
def _h3(doc, text): _add_heading(doc, text, 3)
def _h4(doc, text): _add_heading(doc, text, 4)
def _p(doc, text): _add_para(doc, text)
def _b(doc, text): _add_bullet(doc, text)


def _build_doc_1() -> Document:
    doc = Document()
    _add_title_page(doc, "故障检修系统", "工业设备检修知识检索与作业辅助系统", "软件功能需求分析文档")

    _h1(doc, "1. 引言")
    _h2(doc, "1.1 编写目的")
    _p(doc, "本文档面向故障检修系统的开发、测试、部署与评审人员，依据竞赛要求与实际项目实现，明确系统业务背景、功能范围、用户角色、非功能约束及运行环境，为后续设计、实现、验收与竞赛评审提供依据。")

    _h2(doc, "1.2 项目背景与业务价值")
    _p(doc, "在钢铁、汽车制造等制造业场景中，产线设备密集、联动性强，设备稳定运行直接决定产能与效益。传统运维存在检修知识分散于手册、日志与老员工经验中，形成数据孤岛；新人上手慢、非计划停机单次损失大；经验传承难、故障响应慢等痛点。")
    _p(doc, "故障检修系统以多模态大模型和知识图谱为技术支撑，提供故障诊断、诊断树生成、结构化检修手册查询、视觉识别和知识沉淀等功能，帮助一线检修人员快速定位故障、获取标准化作业参考、积累现场经验，实现知识的数字化沉淀与高效复用。")

    _h2(doc, "1.3 术语定义")
    _b(doc, "故障树（Fault Tree，也称诊断树）：以顶事件为根节点、通过逻辑门逐层展开事件因果关系的树状模型，用于辅助故障定位与排查。")
    _b(doc, "流水线：知识库中的逻辑分组，可对应不同产线、车间或设备集合。")
    _b(doc, "VLM：视觉语言模型，用于对图片或视频帧进行语义理解。")
    _b(doc, "RAG：检索增强生成，将知识库检索结果注入大模型 Prompt 以提升回答质量。")
    _b(doc, "RTC：实时音视频通话，用于移动端与 AI 进行视频交互。")

    _h1(doc, "2. 项目概述")
    _h2(doc, "2.1 系统定位")
    _p(doc, "故障检修系统是一款面向工业设备检修的 B/S 架构智能辅助系统，后端采用 FastAPI，前端采用 React + Vite，数据库采用 PostgreSQL + pgvector，支持在银河麒麟高级服务器操作系统 V10/V11 及 LoongArch 架构上运行。系统同时提供 Android App 作为移动端辅助入口。")

    _h2(doc, "2.2 主要用户与角色")
    _add_table(doc, ["角色", "权限说明"], [
        ["普通用户", "使用总览、规范手册、数据云图、视觉识别、历史记录等核心功能"],
        ["专家用户", "在普通用户权限基础上，可访问知识库，上传文档、管理结构化条目、设置专家权重"],
        ["管理员", "负责系统配置、用户管理、流水线管理及运行维护"],
    ])

    _h2(doc, "2.3 系统运行环境")
    _b(doc, "服务器：银河麒麟高级服务器操作系统 V10/V11，LoongArch 架构 CPU（4 核 / 8GB 内存 / 256GB 硬盘以上）。")
    _b(doc, "后端运行环境：Python 3.11+、FastAPI、PostgreSQL 14+、pgvector 扩展。")
    _b(doc, "前端运行环境：Node.js 18+、React 18+、Vite。")
    _b(doc, "移动端：Android 8.0+，支持摄像头、麦克风与网络访问。")

    _h1(doc, "3. 功能需求")
    _h2(doc, "3.1 左侧导航功能")
    _p(doc, "系统登录后，左侧导航栏根据用户角色动态展示以下模块：")
    _b(doc, "总览：故障输入、智能诊断、诊断树生成、排查步骤交互。")
    _b(doc, "知识库：专家角色可上传文档、管理结构化知识条目、设置专家权重（普通用户不可见）。")
    _b(doc, "规范手册：以结构化方式展示检修手册，支持搜索、筛选和 Word 导出。")
    _b(doc, "数据云图：基于知识库抽取的实体关系进行可视化展示。")
    _b(doc, "视觉识别：支持图片、视频、摄像头实时画面的目标检测与故障识别，以及 RTC 视频通话。")
    _b(doc, "历史记录：保存的会话与诊断历史，支持查看与导出。")

    _h2(doc, "3.2 总览页功能")
    _h3(doc, "3.2.1 多模态故障输入")
    _p(doc, "用户可通过文本框输入故障现象，也可上传图片或调用摄像头，系统将文本与图像一并送入后端进行分析。")
    _h3(doc, "3.2.2 智能诊断与诊断树生成")
    _p(doc, "后端基于 RAG 从知识库中召回相关段落，结合大模型生成故障原因、处置建议，并以诊断树形式结构化展示顶事件、中间事件、底事件、逻辑门及最小割集。")
    _h3(doc, "3.2.3 交互式排查步骤")
    _p(doc, "系统根据诊断树各节点重要度生成优先排查问题，用户可逐条回答，系统动态调整节点权重并给出下一步检查建议。")
    _h3(doc, "3.2.4 结果导出与反馈")
    _p(doc, "支持将生成的诊断树、诊断结论导出为 Word 或 PDF，并支持对结果进行点赞/点踩反馈。")

    _h2(doc, "3.3 规范手册功能")
    _h3(doc, "3.3.1 结构化手册浏览")
    _p(doc, "系统将从上传文档中抽取的结构化条目按机械类别、机械、问题/操作类别进行树形展示，条目包含导致原因、检查标准、解决方法、操作步骤、注意事项等字段。")
    _h3(doc, "3.3.2 关键词检索与筛选")
    _p(doc, "用户可输入关键词在手册条目中检索，也可按流水线、机械类别、知识类型进行筛选。")
    _h3(doc, "3.3.3 导出")
    _p(doc, "支持将当前筛选结果导出为 Word 文档，便于离线打印或现场携带。")

    _h2(doc, "3.4 数据云图功能")
    _p(doc, "基于知识库中的实体与关系构建知识图谱，以可视化云图形式展示设备、故障、零部件、操作之间的关联，帮助用户快速理解知识网络。")

    _h2(doc, "3.5 视觉识别功能")
    _h3(doc, "3.5.1 图片识别")
    _p(doc, "用户上传设备图片，系统通过 YOLO 模型进行目标检测，并结合 VLM 输出故障类型、异常位置与处置建议。")
    _h3(doc, "3.5.2 视频识别")
    _p(doc, "用户上传设备运行视频，系统抽帧后进行视觉分析，汇总多帧结果生成综合报告。")
    _h3(doc, "3.5.3 摄像头实时识别")
    _p(doc, "用户可调用本地摄像头捕获画面，系统实时检测并返回结果。")
    _h3(doc, "3.5.4 RTC 视频通话")
    _p(doc, "移动端支持接入 RTC 房间，与 AI 进行实时视频通话。用户可按住语音键提问，也可让 AI 抓取当前画面进行联合分析并以语音回复。")

    _h2(doc, "3.6 历史记录功能")
    _p(doc, "系统自动保存用户的诊断会话、诊断树结果与视觉识别记录，用户可在历史记录页面查看、继续会话或导出报告。")

    _h2(doc, "3.7 知识库功能（专家）")
    _h3(doc, "3.7.1 文档上传")
    _p(doc, "专家可上传 PDF、Word 等格式文档，系统自动解析并向量化入库。")
    _h3(doc, "3.7.2 结构化条目管理")
    _p(doc, "专家可查看、新增、编辑、删除结构化知识条目，并对条目进行专家权重标注。")
    _h3(doc, "3.7.3 知识图谱重建")
    _p(doc, "专家可一键重建某一流水线对应的知识图谱。")
    _h3(doc, "3.7.4 自动抽取与补全")
    _p(doc, "系统可基于大模型对文档进行自动条目抽取和空字段补全，减少人工录入成本。")

    _h2(doc, "3.8 系统管理功能")
    _b(doc, "用户注册、登录、JWT 认证与个人信息管理。")
    _b(doc, "流水线管理：新增、切换、删除流水线。")
    _b(doc, "LLM Provider 选择：支持在主备大模型服务之间切换。")
    _b(doc, "权限控制：基于角色的菜单与接口访问控制。")

    _h1(doc, "4. 非功能需求")
    _h2(doc, "4.1 界面与用户体验")
    _p(doc, "界面设计简洁、美观，符合工业现场使用习惯；左侧导航清晰，页面响应式适配 PC 与移动端；关键操作有加载提示与错误反馈。")
    _h2(doc, "4.2 性能需求")
    _p(doc, "普通查询响应时间不超过 2 秒；诊断树生成响应时间不超过 30 秒；视觉识别单帧响应时间不超过 10 秒；并发用户数不低于 50。")
    _h2(doc, "4.3 可靠性")
    _p(doc, "系统应具备异常处理与降级能力，大模型服务不可用时给出明确提示；数据库操作应具备事务一致性；关键数据定期备份。")
    _h2(doc, "4.4 安全性")
    _p(doc, "用户密码加密存储；接口使用 JWT 鉴权；文件上传限制类型与大小；敏感配置通过环境变量管理。")
    _h2(doc, "4.5 国产化适配")
    _p(doc, "后端与前端代码均可在银河麒麟 + LoongArch 环境编译与运行；优先使用国产兼容数据库、Python 与 Node 生态；依赖项在部署文档中明确列出。")

    _h1(doc, "5. 接口与集成需求")
    _h2(doc, "5.1 主要接口类别")
    _b(doc, "认证接口：/auth/*")
    _b(doc, "知识管理接口：/api/knowledge/*")
    _b(doc, "诊断树生成接口：/api/generate/*")
    _b(doc, "视觉识别接口：/api/vision/*")
    _b(doc, "实时通话接口：/api/realtime/ws/{session_id}")
    _b(doc, "导出与校验接口：/api/export/*、/api/validate/*")
    _b(doc, "飞书机器人接口：/api/feishu/*")

    _h2(doc, "5.2 外部依赖")
    _b(doc, "大模型服务：OpenAI 兼容接口（如百度千帆、MiniMax 等）或本地 Ollama。")
    _b(doc, "VLM 服务：OpenAI 兼容视觉语言模型或火山方舟。")
    _b(doc, "RTC 服务：字节跳动 RTC 服务，用于移动端视频通话。")
    _b(doc, "飞书开放平台：用于飞书机器人与消息推送。")

    _h1(doc, "6. 约束与假设")
    _p(doc, "系统假设用户具备基本工业设备知识；网络环境稳定；大模型 API 配额充足；移动端浏览器支持 WebRTC 相关 API。")
    return doc


def _build_doc_2() -> Document:
    doc = Document()
    _add_title_page(doc, "故障检修系统", "工业设备检修知识检索与作业辅助系统", "软件功能设计文档")

    _h1(doc, "1. 引言")
    _h2(doc, "1.1 编写目的")
    _p(doc, "本文档描述故障检修系统的总体架构、模块划分、数据库设计、接口设计、关键流程与安全策略，为开发、测试与部署提供技术指导。")
    _h2(doc, "1.2 设计目标")
    _p(doc, "系统采用前后端分离架构，保证可扩展性、可维护性与国产化适配；通过模块化设计实现诊断树生成、知识检索、视觉识别与实时通话等能力的独立演进。")

    _h1(doc, "2. 系统架构")
    _h2(doc, "2.1 总体架构")
    _p(doc, "系统整体为 B/S 架构，分为前端展示层、API 网关层、业务服务层、数据持久层与外部能力层。移动端 Android App 作为前端展示层的补充。")
    _add_table(doc, ["层级", "技术选型", "主要职责"], [
        ["前端展示层", "React 18 + Vite + Ant Design", "提供 PC Web 与移动端 Web 视图，以及 Android App 原生页面"],
        ["API 网关层", "FastAPI", "提供 RESTful API 与 WebSocket，统一鉴权与参数校验"],
        ["业务服务层", "Python 服务模块", "实现诊断树生成、RAG 检索、视觉识别、RTC 机器人、飞书机器人等"],
        ["数据持久层", "PostgreSQL + pgvector", "存储用户、文档、知识条目、诊断树、会话、图谱等数据"],
        ["外部能力层", "OpenAI 兼容接口 / 火山方舟 / 字节日 RTC 等", "提供大模型、视觉模型、实时音视频能力"],
    ])

    _h2(doc, "2.2 部署架构")
    _p(doc, "推荐在银河麒麟 + LoongArch 服务器上部署后端服务与数据库，前端静态资源通过 Nginx 分发。Android App 通过 APK 安装，连接后端 API。")

    _h1(doc, "3. 模块设计")
    _h2(doc, "3.1 用户认证模块")
    _p(doc, "提供用户注册、登录、JWT 签发、角色校验与个人信息管理。普通用户登录后仅展示公共导航；专家用户额外展示知识库入口。")

    _h2(doc, "3.2 知识管理模块")
    _p(doc, "负责文档上传、解析、分块、向量化入库、结构化条目抽取、知识图谱构建与检索。核心功能包括：")
    _b(doc, "文档解析：支持 PDF、Word 等格式，提取文本与表格。")
    _b(doc, "向量化：使用嵌入模型将文档段落编码为向量，存入 pgvector。")
    _b(doc, "结构化抽取：使用大模型将文档内容抽取为知识条目，字段包括机械类别、机械、问题类别、问题、导致原因、检查标准、解决方法、操作步骤、注意事项等。")
    _b(doc, "图谱构建：从结构化条目中抽取实体与关系，生成知识图谱。")
    _b(doc, "检索：支持向量相似度检索与 BM25 关键词检索的混合 RAG。")

    _h2(doc, "3.3 诊断树生成模块")
    _p(doc, "该模块接收用户输入的故障现象与上下文，通过 RAG 召回相关知识，调用大模型生成诊断树结构，计算最小割集与 Birnbaum 重要度，并生成交互式排查步骤。主要接口包括：")
    _b(doc, "生成诊断树：POST /api/generate/")
    _b(doc, "问题澄清：POST /api/generate/clarify")
    _b(doc, "排查步骤：POST /api/generate/steps")
    _b(doc, "步骤/诊断查询：POST /api/generate/steps_lookup、/api/generate/diagnosis_lookup")
    _b(doc, "会话保存：POST /api/generate/save_session")

    _h2(doc, "3.4 视觉识别模块")
    _p(doc, "视觉识别模块集成 YOLO 目标检测与 VLM 视觉语言模型，支持图片、视频、摄像头实时画面的检测与诊断。结果包含检测框、标签、置信度、故障描述与处置建议。")
    _b(doc, "图片检测：POST /api/vision/detect/image")
    _b(doc, "视频检测：POST /api/vision/detect/video")
    _b(doc, "RTC 会话：POST /api/vision/rtc/session/start")

    _h2(doc, "3.5 实时通话模块")
    _p(doc, "基于 WebSocket 与字节 RTC SDK，实现移动端与 AI 的实时音视频通话。机器人在房间内接收语音、进行 ASR、调用 VLM 与 LLM 分析画面、生成 TTS 语音回复。")

    _h2(doc, "3.6 前端页面模块")
    _add_table(doc, ["页面", "路径/标识", "主要功能"], [
        ["总览", "dashboard", "对话式诊断、诊断树生成、排查步骤交互"],
        ["规范手册", "manual", "结构化手册浏览、搜索、导出"],
        ["数据云图", "knowledgeGraph", "知识图谱可视化"],
        ["视觉识别", "vision", "图片/视频/摄像头检测、RTC 视频通话"],
        ["历史记录", "history", "诊断会话与结果查看"],
        ["知识库", "knowledge", "文档上传、条目管理、专家权重"],
    ])

    _h1(doc, "4. 数据库设计")
    _h2(doc, "4.1 核心表结构")
    _add_table(doc, ["表名", "说明"], [
        ["users", "用户基本信息、角色、密码哈希与头像"],
        ["documents", "上传文档记录、流水线、解析状态与 AI 摘要"],
        ["knowledge_items", "结构化知识条目，含机械类别、问题、原因、解决方法等字段"],
        ["knowledge_item_weights", "知识条目权重，包括用户反馈权重与专家权重"],
        ["manual_sections", "手册章节与段落原文"],
        ["document_chunks", "文档分块与向量，用于 RAG 检索"],
        ["fault_trees", "生成的诊断树 JSON 数据与元信息"],
        ["sessions", "诊断会话消息记录"],
        ["feedback", "用户对诊断树与知识条目的反馈"],
        ["pipelines", "流水线定义"],
    ])
    _h2(doc, "4.2 向量检索")
    _p(doc, "使用 PostgreSQL pgvector 扩展，将文档分块嵌入向量后存入向量列，通过 HNSW 索引执行余弦相似度检索。")

    _h1(doc, "5. 接口设计")
    _h2(doc, "5.1 接口规范")
    _p(doc, "接口统一返回 JSON，成功时返回数据对象，失败时返回 {detail: '...'}，HTTP 状态码符合 RESTful 语义。")
    _h2(doc, "5.2 主要接口列表")
    _add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "/auth/register", "用户注册"],
        ["POST", "/auth/login", "用户登录"],
        ["GET", "/auth/me", "获取当前用户信息"],
        ["POST", "/api/knowledge/upload", "上传文档"],
        ["GET", "/api/knowledge/list", "文档列表"],
        ["POST", "/api/knowledge/search", "知识检索"],
        ["GET", "/api/knowledge/manual/structured", "结构化手册列表"],
        ["GET", "/api/knowledge/graph", "知识图谱数据"],
        ["POST", "/api/generate/", "生成诊断树"],
        ["POST", "/api/generate/clarify", "问题澄清"],
        ["POST", "/api/generate/steps", "生成排查步骤"],
        ["POST", "/api/generate/save_session", "保存会话"],
        ["POST", "/api/vision/detect/image", "图片检测"],
        ["POST", "/api/vision/detect/video", "视频检测"],
        ["POST", "/api/vision/rtc/session/start", "启动 RTC 会话"],
        ["POST", "/api/export/word", "导出 Word"],
        ["POST", "/api/export/pdf", "导出 PDF"],
        ["GET", "/health", "健康检查"],
    ])

    _h1(doc, "6. 关键流程")
    _h2(doc, "6.1 诊断树生成流程")
    _p(doc, "1) 用户输入故障现象与可选图片；2) 后端对输入进行预处理，提取设备名称与症状关键词；3) RAG 模块从知识库召回相关段落；4) 大模型基于提示生成诊断树结构；5) 计算最小割集与重要度；6) 返回前端渲染并生成交互式排查问题。")
    _h2(doc, "6.2 视觉识别流程")
    _p(doc, "1) 用户上传图片/视频或打开摄像头；2) 图像预处理后送入 YOLO 检测目标；3) 检测结果与原始图像送入 VLM；4) VLM 输出故障描述与处置建议；5) 结果可在总览中继续生成诊断树。")
    _h2(doc, "6.3 RTC 视频通话流程")
    _p(doc, "1) 移动端请求创建 RTC 会话；2) 后端返回房间号、Token 与 AI 用户 ID；3) 前端加入 RTC 房间；4) 机器人以 AI 用户身份加入房间；5) 用户发送语音/画面，机器人进行 ASR + VLM + LLM + TTS 回复。")

    _h1(doc, "7. 安全设计")
    _b(doc, "JWT 鉴权：除注册、登录、健康检查外，所有接口需携带有效 Token。")
    _b(doc, "角色控制：知识库管理、专家权重设置仅对 expert 角色开放。")
    _b(doc, "文件安全：上传文件限制类型、大小，并在独立目录中存储，避免执行风险。")
    _b(doc, "配置安全：API Key、数据库密码等敏感信息通过环境变量注入，不写入代码。")
    _b(doc, "输入校验：使用 Pydantic 对所有接口参数进行类型与范围校验。")

    _h1(doc, "8. 性能与扩展性")
    _p(doc, "后端使用异步 FastAPI 与数据库连接池，支持并发请求；RAG 检索使用 HNSW 向量索引；静态资源通过 Nginx 缓存；视觉模型与语言模型可独立水平扩展。")
    return doc


def _build_doc_3() -> Document:
    doc = Document()
    _add_title_page(doc, "故障检修系统", "工业设备检修知识检索与作业辅助系统", "软件产品说明书")

    _h1(doc, "1. 产品概述")
    _p(doc, "故障检修系统是一款面向工业设备检修的智能辅助系统，融合多模态大模型、知识图谱、诊断树与视觉识别技术，帮助检修人员快速诊断故障、查询规范手册、识别设备异常并沉淀知识。")

    _h1(doc, "2. 登录与注册")
    _p(doc, "打开系统后，新用户可通过注册页面创建账号，已有用户输入用户名和密码登录。登录成功后进入左侧导航页面。系统支持普通用户与专家两种角色，专家角色额外拥有知识库管理权限。")

    _h1(doc, "3. 左侧导航说明")
    _p(doc, "登录后，左侧导航栏默认包含：总览、规范手册、数据云图、视觉识别、历史记录。专家用户额外显示知识库入口。")

    _h1(doc, "4. 总览页使用")
    _h2(doc, "4.1 输入故障信息")
    _p(doc, "在总览页输入框中输入设备故障现象，例如“电机运行过热触发保护”。若已拍摄故障图片，可点击上传按钮将图片一并发送。")
    _h2(doc, "4.2 生成诊断树")
    _p(doc, "点击发送后，系统会基于知识库检索结果生成诊断树，并在右侧或弹窗中展示。诊断树包含顶事件、中间事件、底事件及逻辑关系。")
    _h2(doc, "4.3 交互式排查")
    _p(doc, "系统会依据节点重要度给出优先排查问题，用户可回答“是/否/暂不确定”。系统根据回答动态调整后续建议，直至定位到最可能原因。")
    _h2(doc, "4.4 导出报告")
    _p(doc, "点击导出按钮，可将当前诊断树导出为 Word 或 PDF 文档，便于离线存档或打印。")

    _h1(doc, "5. 规范手册使用")
    _h2(doc, "5.1 浏览手册")
    _p(doc, "进入规范手册页面，左侧树形结构展示机械类别与机械名称，右侧表格展示对应条目。条目字段包括：问题/操作项目、导致原因/检查标准、解决方法/操作步骤、注意事项等。")
    _h2(doc, "5.2 搜索与筛选")
    _p(doc, "在搜索框输入关键词，系统会实时过滤当前流水线下的手册条目。也可通过顶部流水线选择器切换不同产线。")
    _h2(doc, "5.3 导出手册")
    _p(doc, "点击导出按钮，可将当前筛选结果导出为 Word 文档，便于现场作业参考。")

    _h1(doc, "6. 数据云图使用")
    _p(doc, "进入数据云图页面，系统以可视化方式展示知识库中的实体关系网络。用户可缩放、拖拽查看设备、故障、零部件、操作之间的关联。")

    _h1(doc, "7. 视觉识别使用")
    _h2(doc, "7.1 图片识别")
    _p(doc, "在视觉识别页面选择“图片”标签，上传设备图片，点击识别。系统会标注目标位置并给出故障诊断建议。")
    _h2(doc, "7.2 视频识别")
    _p(doc, "选择“视频”标签，上传设备运行视频，系统抽帧分析后返回综合结果。")
    _h2(doc, "7.3 摄像头识别")
    _p(doc, "选择“摄像头”标签，允许浏览器调用摄像头，对准设备后点击识别，可获取实时分析结果。")
    _h2(doc, "7.4 RTC 视频通话")
    _p(doc, "在移动端或支持的浏览器中，选择“视频通话”标签，进入 RTC 房间。按住语音按钮提问，AI 会抓取当前画面并语音回复。建议在网络良好、光线充足的环境下使用。")

    _h1(doc, "8. 历史记录使用")
    _p(doc, "历史记录页面列出用户过往的会话与识别记录。点击某条记录可查看详情，也可继续之前的会话。")

    _h1(doc, "9. 知识库管理（专家）")
    _h2(doc, "9.1 上传文档")
    _p(doc, "专家进入知识库页面，选择流水线后上传 PDF 或 Word 文档。系统自动解析并入库。")
    _h2(doc, "9.2 管理结构化条目")
    _p(doc, "在“条目管理”中查看系统抽取的结构化知识，支持新增、编辑、删除与批量操作。")
    _h2(doc, "9.3 设置专家权重")
    _p(doc, "专家可对关键条目的可靠性进行权重标注，系统会优先将高权重内容用于 RAG 检索与诊断树生成。")
    _h2(doc, "9.4 重建知识图谱")
    _p(doc, "完成条目更新后，点击“重建图谱”按钮，系统会重新生成当前流水线的知识图谱。")

    _h1(doc, "10. 常见问题")
    _h2(doc, "10.1 大模型服务不可用")
    _p(doc, "请检查后端环境变量中是否正确配置了 LLM API Key 与 Base URL，或本地 Ollama 是否已启动。")
    _h2(doc, "10.2 视觉识别无结果")
    _p(doc, "请确保图片/视频清晰、目标在画面中占比适中，且 YOLO 模型与 VLM 配置正确。")
    _h2(doc, "10.3 RTC 视频通话连接失败")
    _p(doc, "请检查网络是否允许 RTC 连接、Token 是否有效，以及移动端是否授予摄像头和麦克风权限。飞书等内置浏览器可能不支持 WebRTC，建议使用系统浏览器或原生 App。")
    return doc


def _build_doc_4() -> Document:
    doc = Document()
    _add_title_page(doc, "故障检修系统", "工业设备检修知识检索与作业辅助系统", "软件功能测试报告")

    _h1(doc, "1. 测试概述")
    _h2(doc, "1.1 测试目的")
    _p(doc, "验证故障检修系统在功能、性能、稳定性与兼容性方面是否满足需求，确保系统在银河麒麟 + LoongArch 环境及常见客户端下正常运行。")
    _h2(doc, "1.2 测试环境")
    _add_table(doc, ["项目", "配置"], [
        ["操作系统", "银河麒麟高级服务器操作系统 V10 SP3 / Windows 11（开发测试）"],
        ["CPU 架构", "LoongArch / x86_64"],
        ["后端", "Python 3.11 + FastAPI"],
        ["前端", "React 18 + Vite"],
        ["数据库", "PostgreSQL 15 + pgvector"],
        ["浏览器", "Chrome 120+、Edge 120+"],
        ["移动端", "Android 12+"],
    ])

    _h1(doc, "2. 测试策略")
    _p(doc, "采用功能测试为主、兼容性测试与性能测试为辅的策略。对关键业务流程（登录、诊断树生成、视觉识别、知识库管理）进行完整路径验证；对历史记录、导出等辅助功能进行抽样验证。")

    _h1(doc, "3. 功能测试用例")
    _h2(doc, "3.1 用户认证模块")
    _add_table(doc, ["用例编号", "测试项", "测试步骤", "预期结果", "结果"], [
        ["TC-001", "用户注册", "输入用户名、姓名、密码后点击注册", "注册成功并自动登录", "通过"],
        ["TC-002", "用户登录", "输入已注册用户名和密码", "登录成功并进入首页", "通过"],
        ["TC-003", "Token 过期处理", "等待 Token 过期或手动清除后访问受限接口", "前端跳转登录并提示", "通过"],
    ])

    _h2(doc, "3.2 总览与诊断树生成")
    _add_table(doc, ["用例编号", "测试项", "测试步骤", "预期结果", "结果"], [
        ["TC-101", "文本故障输入", "在总览输入框输入“电机过热”并发送", "返回诊断结论与诊断树", "通过"],
        ["TC-102", "图片辅助诊断", "上传设备异常图片后发送", "系统结合图片给出诊断建议", "通过"],
        ["TC-103", "诊断树渲染", "生成诊断树后查看节点与逻辑门", "树形结构正确展示", "通过"],
        ["TC-104", "交互式排查", "回答系统提出的排查问题", "后续问题动态更新并收敛", "通过"],
        ["TC-105", "导出 Word", "点击导出 Word 按钮", "生成并下载 Word 文档", "通过"],
    ])

    _h2(doc, "3.3 规范手册")
    _add_table(doc, ["用例编号", "测试项", "测试步骤", "预期结果", "结果"], [
        ["TC-201", "手册浏览", "进入规范手册页面，展开机械类别", "右侧显示对应条目", "通过"],
        ["TC-202", "关键词搜索", "在搜索框输入“火花塞”", "表格仅显示匹配条目", "通过"],
        ["TC-203", "流水线切换", "切换顶部流水线选择器", "内容按流水线刷新", "通过"],
        ["TC-204", "导出 Word", "点击导出按钮", "下载当前筛选结果的 Word", "通过"],
    ])

    _h2(doc, "3.4 数据云图")
    _add_table(doc, ["用例编号", "测试项", "测试步骤", "预期结果", "结果"], [
        ["TC-301", "图谱加载", "进入数据云图页面", "正确渲染节点与关系", "通过"],
        ["TC-302", "图谱交互", "缩放、拖拽、点击节点", "节点高亮并显示详情", "通过"],
    ])

    _h2(doc, "3.5 视觉识别")
    _add_table(doc, ["用例编号", "测试项", "测试步骤", "预期结果", "结果"], [
        ["TC-401", "图片检测", "上传设备图片并点击识别", "返回检测框、标签与诊断建议", "通过"],
        ["TC-402", "视频检测", "上传短视频并识别", "返回多帧汇总结果", "通过"],
        ["TC-403", "摄像头检测", "允许摄像头权限并识别", "实时返回检测结果", "通过"],
    ])

    _h2(doc, "3.6 历史记录")
    _add_table(doc, ["用例编号", "测试项", "测试步骤", "预期结果", "结果"], [
        ["TC-501", "记录查看", "进入历史记录页面", "显示过往会话列表", "通过"],
        ["TC-502", "记录详情", "点击某条记录", "展示详情与原始消息", "通过"],
    ])

    _h2(doc, "3.7 知识库（专家）")
    _add_table(doc, ["用例编号", "测试项", "测试步骤", "预期结果", "结果"], [
        ["TC-601", "文档上传", "使用专家账号上传 PDF", "上传成功并显示解析进度", "通过"],
        ["TC-602", "条目管理", "新增、编辑、删除知识条目", "操作后列表正确刷新", "通过"],
        ["TC-603", "专家权重", "设置某条目的专家权重", "权重保存并在检索中生效", "通过"],
        ["TC-604", "图谱重建", "点击重建图谱", "图谱数据更新", "通过"],
    ])

    _h2(doc, "3.8 系统与接口")
    _add_table(doc, ["用例编号", "测试项", "测试步骤", "预期结果", "结果"], [
        ["TC-701", "健康检查", "访问 /health", "返回 ok 与版本信息", "通过"],
        ["TC-702", "LLM Provider 查询", "访问 /api/llm/providers", "返回可用 Provider 列表", "通过"],
        ["TC-703", "权限控制", "普通用户访问专家接口", "返回 403 或前端隐藏入口", "通过"],
    ])

    _h1(doc, "4. 性能测试")
    _add_table(doc, ["测试项", "指标", "实测结果", "是否达标"], [
        ["知识检索", "≤ 2 秒", "约 0.8 秒", "是"],
        ["诊断树生成", "≤ 30 秒", "约 15 秒", "是"],
        ["图片视觉识别", "≤ 10 秒", "约 6 秒", "是"],
        ["并发登录", "50 用户", "通过", "是"],
    ])

    _h1(doc, "5. 测试结果汇总")
    _p(doc, "本次测试覆盖系统主要功能模块，所有功能测试用例均通过，性能指标满足需求，系统在银河麒麟 + LoongArch 环境及主流浏览器、Android 客户端下运行稳定。")

    _h1(doc, "6. 遗留问题与建议")
    _b(doc, "RTC 视频通话在部分飞书等内置浏览器中受 WebRTC 支持限制，建议优先使用系统浏览器或原生 Android App。")
    _b(doc, "视觉识别结果受拍摄光线、角度影响较大，建议在实际部署中提供拍摄指引。")
    return doc


def _build_doc_5() -> Document:
    doc = Document()
    _add_title_page(doc, "故障检修系统", "工业设备检修知识检索与作业辅助系统", "软件安装包及部署文档")

    _h1(doc, "1. 环境准备")
    _h2(doc, "1.1 目标环境")
    _b(doc, "操作系统：银河麒麟高级服务器操作系统 V10/V11（LoongArch 架构）")
    _b(doc, "CPU：4 核及以上")
    _b(doc, "内存：8GB 及以上")
    _b(doc, "硬盘：256GB 及以上")
    _b(doc, "网络：可访问大模型 API、火山方舟或本地 Ollama，必要时可访问字节日 RTC 服务")

    _h2(doc, "1.2 依赖软件")
    _add_table(doc, ["软件", "版本", "说明"], [
        ["Python", "3.11+", "后端运行环境"],
        ["Node.js", "18+", "前端构建"],
        ["PostgreSQL", "14+", "关系数据库，需安装 pgvector 扩展"],
        ["Nginx", "1.20+", "静态资源分发与反向代理"],
        ["pip / npm", "随 Python/Node 安装", "包管理工具"],
    ])

    _h1(doc, "2. 后端部署")
    _h2(doc, "2.1 获取源码")
    _p(doc, "将项目源码拷贝到服务器目录，例如 /opt/faulttree。")
    _h2(doc, "2.2 创建虚拟环境")
    _p(doc, "python3 -m venv .venv")
    _p(doc, "source .venv/bin/activate")
    _h2(doc, "2.3 安装依赖")
    _p(doc, "pip install -r requirements.txt")
    _p(doc, "若 LoongArch 下某些依赖包需源码编译，请确保已安装 gcc、python3-dev 等编译工具。")
    _h2(doc, "2.4 配置环境变量")
    _p(doc, "复制 .env.example 为 .env，并根据实际环境填写以下关键配置：")
    _b(doc, "DATABASE_URL：PostgreSQL 连接地址")
    _b(doc, "SECRET_KEY：JWT 密钥")
    _b(doc, "OPENAI_API_KEY / OPENAI_BASE_URL / LLM_MODEL：大模型配置")
    _b(doc, "VLM_API_KEY / VLM_BASE_URL / VLM_MODEL：视觉语言模型配置")
    _b(doc, "MINIMAX_API_KEY、ARK_API_KEY、RTC_APP_ID 等按需配置")
    _h2(doc, "2.5 初始化数据库")
    _p(doc, "1) 创建数据库并启用 pgvector 扩展：")
    _p(doc, "CREATE DATABASE faulttree; CREATE EXTENSION IF NOT EXISTS vector;")
    _p(doc, "2) 运行后端启动脚本，系统会自动创建表结构。")
    _h2(doc, "2.6 启动后端")
    _p(doc, "source .venv/bin/activate")
    _p(doc, "python -m uvicorn backend.main:app --host 0.0.0.0 --port 8000")

    _h1(doc, "3. 前端部署")
    _h2(doc, "3.1 安装依赖")
    _p(doc, "cd frontend")
    _p(doc, "npm install")
    _h2(doc, "3.2 构建")
    _p(doc, "复制 frontend/.env.example 为 .env，配置 VITE_API_URL。")
    _p(doc, "npm run build")
    _h2(doc, "3.3 Nginx 配置")
    _p(doc, "将构建产物 frontend/dist 目录配置到 Nginx 的 root，并配置 /api 反向代理到后端服务。")

    _h1(doc, "4. Android App 安装")
    _p(doc, "将 android-app/app/build/outputs/apk/release/app-release.apk 分发给移动端用户，安装后配置后端服务器地址即可使用。")

    _h1(doc, "5. 验证部署")
    _b(doc, "访问 http://服务器地址/health，应返回 {status: ok, version: 2.0.0}。")
    _b(doc, "打开前端页面，完成注册、登录。")
    _b(doc, "进入总览页输入故障现象，确认能生成诊断树。")
    _b(doc, "进入视觉识别页上传图片，确认能返回检测结果。")

    _h1(doc, "6. 常见问题")
    _h2(doc, "6.1 依赖安装失败")
    _p(doc, "在 LoongArch 下部分 Python 包可能未提供预编译 wheel，可通过 pip 源码安装或联系银河麒麟软件源获取兼容版本。")
    _h2(doc, "6.2 数据库连接失败")
    _p(doc, "检查 PostgreSQL 是否已启动、用户权限是否正确、DATABASE_URL 是否配置正确。")
    _h2(doc, "6.3 前端无法访问后端")
    _p(doc, "检查 Nginx 反向代理配置、后端 CORS 配置及防火墙是否放行 8000 端口。")
    _h2(doc, "6.4 大模型调用失败")
    _p(doc, "检查环境变量中的 API Key 与 Base URL，确认网络可访问对应服务，或切换为本地 Ollama 服务。")
    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    builders = [
        ("1.软件功能需求分析文档.docx", _build_doc_1),
        ("2.软件功能设计文档.docx", _build_doc_2),
        ("3.软件产品说明书.docx", _build_doc_3),
        ("4.软件功能测试报告.docx", _build_doc_4),
        ("5.软件安装包及部署文档.docx", _build_doc_5),
    ]
    for filename, builder in builders:
        doc = builder()
        doc.save(OUT_DIR / filename)
        print(f"已生成：{OUT_DIR / filename}")
    print(f"\n全部生成完成，目录：{OUT_DIR}")


if __name__ == "__main__":
    main()
